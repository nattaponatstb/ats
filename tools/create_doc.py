import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Color palette (Dummy for compatibility) ───────────────────────────────────
NAVY      = RGBColor(0x00, 0x00, 0x00)
BLUE_MED  = RGBColor(0x00, 0x00, 0x00)
BLUE_LT   = RGBColor(0x00, 0x00, 0x00)
GOLD      = RGBColor(0x00, 0x00, 0x00)
GREEN     = RGBColor(0x00, 0x00, 0x00)
GREEN_LT  = RGBColor(0x00, 0x00, 0x00)
RED_D     = RGBColor(0x00, 0x00, 0x00)
RED_LT    = RGBColor(0x00, 0x00, 0x00)
ORANGE    = RGBColor(0x00, 0x00, 0x00)
ORANGE_LT = RGBColor(0x00, 0x00, 0x00)
GRAY_LT   = RGBColor(0x00, 0x00, 0x00)
GRAY_MED  = RGBColor(0x00, 0x00, 0x00)
WHITE     = RGBColor(0xff, 0xff, 0xff)
BLACK     = RGBColor(0x00, 0x00, 0x00)

FONT_TH = 'TH Sarabun New'

# ── Helper Styling Functions ──
def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # thin border
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # black
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_bg(cell, color: RGBColor):
    pass

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    pass

# ── Helper: paragraph ─────────────────────────────────────────────────────────
def add_para(doc_or_cell, text='', size=14, bold=False, italic=False,
             color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, font_name=FONT_TH):
    if hasattr(doc_or_cell, 'paragraphs') and hasattr(doc_or_cell, '_tc'):
        # it's a cell
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
        p.clear()
    else:
        p = doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name      = font_name
        run.font.size      = Pt(size)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) # strictly black text
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_run(para, text, size=14, bold=False, italic=False,
            color=BLACK, font_name=FONT_TH):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name      = font_name
    run.font.size      = Pt(size)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) # strictly black text
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run

# ── Helper: section heading ───────────────────────────────────────────────────
def section_heading(doc, num, title, subtitle=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r1 = p.add_run(f'{num}. {title}')
    r1.bold = True
    r1.font.name = FONT_TH
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.keep_with_next = True
        r2 = p2.add_run(f"({subtitle})")
        r2.italic = True
        r2.font.name = FONT_TH
        r2.font.size = Pt(13)
        r2.font.color.rgb = RGBColor(0, 0, 0)
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

# ── Helper: info box ──────────────────────────────────────────────────────────
def info_box(doc, title, items, bg=None, title_bg=None):
    tbl = doc.add_table(rows=1+len(items), cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)
    
    # title row
    tc = tbl.cell(0,0)
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'  {title}')
    r.bold = True
    r.font.name = FONT_TH
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    for i, item in enumerate(items):
        cell = tbl.cell(i+1, 0)
        p2 = cell.paragraphs[0]
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after  = Pt(3)
        r2 = p2.add_run(f'   {item}')
        r2.font.name = FONT_TH
        r2.font.size = Pt(13)
        r2.font.color.rgb = RGBColor(0, 0, 0)
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE (Plain)
# ══════════════════════════════════════════════════════════════════════════════
# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(40)
p_title.paragraph_format.space_after = Pt(10)
r_t = p_title.add_run("ระบบตารางสอนประจำสัปดาห์")
r_t.bold = True
r_t.font.name = FONT_TH
r_t.font.size = Pt(20)
r_t.font.color.rgb = RGBColor(0, 0, 0)
r_t._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(10)
r_sub = p_sub.add_run("Weekly Timetable Management System")
r_sub.font.name = FONT_TH
r_sub.font.size = Pt(15)
r_sub.font.color.rgb = RGBColor(0, 0, 0)
r_sub._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

# Organization
p_org = doc.add_paragraph()
p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org.paragraph_format.space_before = Pt(0)
p_org.paragraph_format.space_after = Pt(15)
r_org = p_org.add_run("โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก")
r_org.font.name = FONT_TH
r_org.font.size = Pt(15)
r_org.font.color.rgb = RGBColor(0, 0, 0)
r_org._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

# Badge / Spec info
p_badge = doc.add_paragraph()
p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_badge.paragraph_format.space_before = Pt(0)
p_badge.paragraph_format.space_after = Pt(40)
r_b = p_badge.add_run("🌐 Web Application  |  ☁️ Cloud Database  |  📡 Real-Time Sync")
r_b.font.name = FONT_TH
r_b.font.size = Pt(12)
r_b.font.color.rgb = RGBColor(0, 0, 0)
r_b._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

# Bottom info table
tbl_bot = doc.add_table(rows=1, cols=3)
tbl_bot.style = 'Table Grid'
tbl_bot.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tbl_bot)

labels = ['URL', 'ฐานข้อมูล', 'Hosting']
values = ['nattaponatstb.github.io/ats/', 'Firebase Realtime DB', 'GitHub Pages']
for i, (lbl, val) in enumerate(zip(labels, values)):
    c = tbl_bot.cell(0, i)
    p_lbl = c.paragraphs[0]
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lbl.paragraph_format.space_before = Pt(4)
    p_lbl.paragraph_format.space_after  = Pt(0)
    r_lbl = p_lbl.add_run(lbl)
    r_lbl.bold = True
    r_lbl.font.name = FONT_TH
    r_lbl.font.size = Pt(11)
    r_lbl.font.color.rgb = RGBColor(0, 0, 0)
    r_lbl._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    p_val = c.add_paragraph()
    p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_val.paragraph_format.space_before = Pt(0)
    p_val.paragraph_format.space_after  = Pt(4)
    r_val = p_val.add_run(val)
    r_val.font.name = FONT_TH
    r_val.font.size = Pt(11)
    r_val.font.color.rgb = RGBColor(0, 0, 0)
    r_val._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

# page break
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1: ภาพรวมโปรแกรม
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '01', 'ภาพรวมโปรแกรม', 'Program Overview')

add_para(doc,
    'ระบบตารางสอนประจำสัปดาห์ คือ เว็บแอปพลิเคชันสำหรับจัดการตารางสอน '
    'ของโรงเรียนทหารขนส่ง กรมการขนส่งทหารบก  '
    'พัฒนาในรูปแบบ Single Page Application (SPA) '
    'ทำงานผ่าน Browser ทุกเครื่อง ไม่ต้องติดตั้งโปรแกรมใดๆ '
    'และข้อมูลซิงค์แบบ Real-Time ระหว่างผู้ใช้ทุกคนพร้อมกัน',
    size=14, space_after=10)

# 4 highlight boxes
tbl_hl = doc.add_table(rows=2, cols=2)
tbl_hl.style = 'Table Grid'
tbl_hl.alignment = WD_TABLE_ALIGNMENT.CENTER
highlights = [
    ('🖥️', 'ประเภทโปรแกรม',  'Web Application (SPA)\nทำงานบน Browser ทุกระบบ', BLUE_LT,   BLUE_MED),
    ('☁️', 'ฐานข้อมูล',      'Firebase Realtime Database\n(Cloud — Google)', GREEN_LT, GREEN),
    ('📡', 'การซิงค์ข้อมูล', 'Real-Time ทุกคนเห็นข้อมูล\nเดียวกันพร้อมกันทันที', ORANGE_LT, ORANGE),
    ('🔒', 'ระบบสิทธิ์',     'Role-Based Access Control\nSuperAdmin / Admin / Public', RED_LT,  RED_D),
]
positions = [(0,0),(0,1),(1,0),(1,1)]
for (r,c), (icon, title, desc, bg, title_c) in zip(positions, highlights):
    cell = tbl_hl.cell(r, c)
    set_cell_bg(cell, bg)
    p_i = cell.paragraphs[0]
    p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_i.paragraph_format.space_before = Pt(8)
    p_i.paragraph_format.space_after  = Pt(2)
    ri = p_i.add_run(icon)
    ri.font.size = Pt(22)
    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(2)
    rt = p_title.add_run(title)
    rt.bold = True
    rt.font.name = FONT_TH
    rt.font.size = Pt(13)
    rt.font.color.rgb = title_c
    rt._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.space_after  = Pt(8)
    rd = p_desc.add_run(desc)
    rd.font.name = FONT_TH
    rd.font.size = Pt(12)
    rd.font.color.rgb = BLACK
    rd._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_paragraph()

# คุณสมบัติหลัก
add_para(doc, '✦  คุณสมบัติหลักของระบบ', size=15, bold=True, color=NAVY, space_before=6, space_after=4)

features = [
    ('📋', 'จัดการหลักสูตรและตารางสอน', 'สร้าง แก้ไข ลบ หลักสูตร พร้อมรายการสอนประจำสัปดาห์'),
    ('👥', 'ซิงค์แบบ Real-Time', 'ผู้ใช้ทุกคนเห็นการเปลี่ยนแปลงพร้อมกันโดยไม่ต้อง Refresh'),
    ('🖨️', 'พิมพ์ตาราง A4', 'พิมพ์ตารางสอน A4 Landscape ด้วยฟอนต์ Sarabun เหมือนกันทุกเครื่อง'),
    ('📊', 'Export Excel', 'ส่งออกสรุปชั่วโมง รายวิชา รายครู รายสัปดาห์ เป็นไฟล์ Excel'),
    ('🔐', 'ระบบ Login', 'ควบคุมสิทธิ์การเข้าถึงด้วย Username / Password'),
    ('🌐', 'Public Mode', 'เปิดให้บุคคลภายนอกดูตารางได้โดยไม่ต้อง Login'),
    ('💾', 'สำรองข้อมูล', 'Backup อัตโนมัติขึ้น GitHub Pages โดย Super Admin'),
]

tbl_feat = doc.add_table(rows=len(features), cols=3)
tbl_feat.style = 'Table Grid'
tbl_feat.alignment = WD_TABLE_ALIGNMENT.CENTER
col_w = [Cm(1.2), Cm(5.5), Cm(9.8)]
for i, (icon, title, desc) in enumerate(features):
    cells = tbl_feat.row_cells(i)
    bg = GRAY_LT if i % 2 == 0 else WHITE
    for cell in cells:
        set_cell_bg(cell, bg)
    # icon
    p0 = cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after  = Pt(4)
    r0 = p0.add_run(icon)
    r0.font.size = Pt(14)
    # title
    p1 = cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(4)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.name = FONT_TH
    r1.font.size = Pt(13)
    r1.font.color.rgb = NAVY
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    # desc
    p2 = cells[2].paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    r2 = p2.add_run(desc)
    r2.font.name = FONT_TH
    r2.font.size = Pt(12)
    r2.font.color.rgb = BLACK
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2: โครงสร้างระบบ
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '02', 'โครงสร้างระบบ', 'System Architecture')

add_para(doc, 'ระบบประกอบด้วย 3 ชั้นหลัก ทำงานร่วมกันโดยไม่มี Server กลางของตนเอง',
    size=14, space_after=10)

# Architecture diagram (table-based)
layers = [
    ('🖥️  ชั้นที่ 1 — Browser (ผู้ใช้)',
     NAVY, WHITE,
     [
       'ไฟล์หลัก: timetable_military.html  (ไฟล์เดียว ~4,800 บรรทัด)',
       'ประกอบด้วย: HTML (หน้าจอ) + CSS (สไตล์) + JavaScript (ตรรกะทั้งหมด)',
       'localStorage: เก็บข้อมูลชั่วคราวในเครื่อง (Cache) เพื่อความเร็ว',
       'ทำงานบน Chrome / Edge / Firefox ทุกระบบปฏิบัติการ',
     ], BLUE_LT),
    ('☁️  ชั้นที่ 2 — Firebase Realtime Database (ฐานข้อมูลกลาง)',
     GREEN, WHITE,
     [
       'บริการของ Google — ไม่ต้องดูแล Server เอง',
       '/data/courses  — ข้อมูลหลักสูตรทั้งหมด',
       '/data/entries  — รายการสอนทุกรายการ',
       '/data/settings — ตั้งค่าโรงเรียน ครู วิชา',
       '/data/users    — ข้อมูล user + password',
       'รับส่งผ่าน REST API (HTTPS) และ SSE (Real-Time Stream)',
     ], GREEN_LT),
    ('🌐  ชั้นที่ 3 — GitHub Pages (Hosting + สำรองข้อมูล)',
     RGBColor(0x33,0x33,0x33), WHITE,
     [
       'Host ไฟล์ HTML ให้เข้าถึงได้จากทุกที่ผ่าน Internet',
       'URL: https://nattaponatstb.github.io/ats/',
       'เก็บ data.json — snapshot สำรองข้อมูลล่าสุดโดย Super Admin',
       'ถ้า Firebase มีปัญหา ระบบดึงข้อมูลจาก data.json แทน (Fallback)',
     ], GRAY_LT),
]

for layer_title, title_bg, title_fg, items, item_bg in layers:
    tbl_layer = doc.add_table(rows=1+len(items), cols=1)
    tbl_layer.style = 'Table Grid'
    tbl_layer.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hcell = tbl_layer.cell(0,0)
    set_cell_bg(hcell, title_bg)
    ph = hcell.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ph.paragraph_format.space_before = Pt(5)
    ph.paragraph_format.space_after  = Pt(5)
    rh = ph.add_run(f'  {layer_title}')
    rh.bold = True
    rh.font.name = FONT_TH
    rh.font.size = Pt(14)
    rh.font.color.rgb = title_fg
    rh._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    for i, item in enumerate(items):
        ic = tbl_layer.cell(i+1,0)
        set_cell_bg(ic, item_bg)
        pi = ic.paragraphs[0]
        pi.paragraph_format.space_before = Pt(3)
        pi.paragraph_format.space_after  = Pt(3)
        ri = pi.add_run(f'      ▸  {item}')
        ri.font.name = FONT_TH
        ri.font.size = Pt(13)
        ri.font.color.rgb = BLACK
        ri._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    doc.add_paragraph()

# Data flow
add_para(doc, '✦  การไหลของข้อมูล (Data Flow)', size=15, bold=True, color=NAVY, space_before=6, space_after=4)

flow_steps = [
    ('1', 'ผู้ใช้กระทำ (Add / Edit / Delete)',
     'ผู้ใช้กดปุ่มเพิ่ม แก้ไข หรือลบข้อมูล', NAVY),
    ('2', 'บันทึก localStorage ทันที',
     'ข้อมูลถูกบันทึกใน Browser ก่อนเพื่อความเร็ว UI ไม่ค้าง', BLUE_MED),
    ('3', 'Push ขึ้น Firebase (Async)',
     'ระบบส่งเฉพาะ node ที่เปลี่ยนแปลงขึ้น Firebase ในพื้นหลัง', GREEN),
    ('4', 'Firebase ส่ง SSE Event',
     'Firebase แจ้ง Browser ทุกเครื่องที่เชื่อมอยู่ว่ามีการเปลี่ยนแปลง', ORANGE),
    ('5', 'ทุก Browser อัพเดตพร้อมกัน',
     'แต่ละ Browser รับ event อัพเดต localStorage และ re-render UI ทันที', GREEN),
]
tbl_flow = doc.add_table(rows=len(flow_steps), cols=3)
tbl_flow.style = 'Table Grid'
for i, (num, title, desc, color) in enumerate(flow_steps):
    cells = tbl_flow.row_cells(i)
    bg = BLUE_LT if i % 2 == 0 else WHITE
    set_cell_bg(cells[0], color)
    set_cell_bg(cells[1], bg)
    set_cell_bg(cells[2], bg)
    # number
    p0 = cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after  = Pt(6)
    r0 = p0.add_run(num)
    r0.bold = True
    r0.font.name = FONT_TH
    r0.font.size = Pt(16)
    r0.font.color.rgb = WHITE
    r0._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    # title
    p1 = cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(6)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.name = FONT_TH
    r1.font.size = Pt(13)
    r1.font.color.rgb = NAVY
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    # desc
    p2 = cells[2].paragraphs[0]
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after  = Pt(6)
    r2 = p2.add_run(desc)
    r2.font.name = FONT_TH
    r2.font.size = Pt(12)
    r2.font.color.rgb = BLACK
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3: หน้าจอและส่วนต่างๆ
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '03', 'หน้าจอและส่วนการทำงาน', 'Screens & Modules')

screens = [
    ('🔐', 'หน้า Login',       NAVY,    WHITE,
     'ผู้ใช้กรอก Username และ Password\nระบบดึงข้อมูล user จาก Firebase ก่อน Login ทุกครั้ง\nเพื่อให้ใช้งานได้ทุกเครื่องทุกที่'),
    ('📅', 'ตารางสอน (หลัก)', BLUE_MED, WHITE,
     'แสดงรายการสอนประจำสัปดาห์\nเพิ่ม / แก้ไข / ลบ / คัดลอก รายการสอน\nกำหนดวันหยุดประจำสัปดาห์'),
    ('📊', 'สรุปชั่วโมง',      GREEN,   WHITE,
     'สรุปชั่วโมงสอนรายวิชา / รายครู / รายสัปดาห์\nExport เป็น Excel สำหรับนำไปรายงาน'),
    ('🖨️', 'พิมพ์ตาราง',      ORANGE,  WHITE,
     'พิมพ์ตาราง A4 Landscape ได้ทุกเครื่อง\nฟอนต์ Sarabun จาก Google Fonts (เหมือนกันทุกที่)\nบันทึกเป็น PDF ได้'),
    ('⚙️', 'ตั้งค่า',          RED_D,   WHITE,
     'ตั้งค่าชื่อโรงเรียน / ชื่อย่อ\nจัดการรายชื่อครู วิชา ประเภทหลักสูตร สถานที่\nตั้งค่า GitHub Token สำหรับ backup'),
    ('👁️', 'Public Mode',      RGBColor(0x44,0x44,0x44), WHITE,
     'ดูตารางสอนโดยไม่ต้อง Login\nเฉพาะหลักสูตรที่ Super Admin เปิด Public\nไม่สามารถแก้ไขข้อมูลได้'),
]

tbl_scr = doc.add_table(rows=3, cols=2)
tbl_scr.style = 'Table Grid'
for idx, (icon, name, bg, fg, desc) in enumerate(screens):
    r, c = divmod(idx, 2)
    cell = tbl_scr.cell(r, c)
    set_cell_bg(cell, RGBColor(0xf8,0xf9,0xfc))
    # header bar
    p_hdr = cell.paragraphs[0]
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(0)
    # insert colored bar via separate paragraph approach
    # Use run color for title
    r_icon2 = p_hdr.add_run(f'{icon} ')
    r_icon2.font.size = Pt(14)
    r_name = p_hdr.add_run(name)
    r_name.bold = True
    r_name.font.name = FONT_TH
    r_name.font.size = Pt(14)
    r_name.font.color.rgb = bg
    r_name._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    p_hdr.paragraph_format.space_before = Pt(8)
    p_desc2 = cell.add_paragraph()
    p_desc2.paragraph_format.space_before = Pt(2)
    p_desc2.paragraph_format.space_after  = Pt(8)
    r_desc2 = p_desc2.add_run(desc)
    r_desc2.font.name = FONT_TH
    r_desc2.font.size = Pt(12)
    r_desc2.font.color.rgb = RGBColor(0x33,0x33,0x33)
    r_desc2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4: กลุ่มผู้ใช้งาน
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '04', 'กลุ่มผู้ใช้งานและสิทธิ์', 'User Groups & Permissions')

add_para(doc, 'ระบบแบ่งผู้ใช้งานออกเป็น 3 กลุ่ม แต่ละกลุ่มมีสิทธิ์การเข้าถึงที่แตกต่างกัน',
    size=14, space_after=10)

# --- Super Admin ---
tbl_sa = doc.add_table(rows=1, cols=1)
tbl_sa.style = 'Table Grid'
tbl_sa.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_sa = tbl_sa.cell(0,0)
set_cell_bg(cell_sa, NAVY)
set_cell_border(cell_sa,
    left={'val':'single','sz':16,'color':'C89F2E'},
    top={'val':'single','sz':4,'color':'444444'},
    bottom={'val':'single','sz':4,'color':'444444'},
    right={'val':'single','sz':4,'color':'444444'})
p_sa = cell_sa.paragraphs[0]
p_sa.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sa.paragraph_format.space_before = Pt(8)
p_sa.paragraph_format.space_after  = Pt(4)
r_sa = p_sa.add_run('  👑  Super Admin  —  ผู้ดูแลระบบสูงสุด')
r_sa.bold = True
r_sa.font.name = FONT_TH
r_sa.font.size = Pt(16)
r_sa.font.color.rgb = WHITE
r_sa._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
p_sa2 = cell_sa.add_paragraph()
p_sa2.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sa2.paragraph_format.space_before = Pt(0)
p_sa2.paragraph_format.space_after  = Pt(8)
r_sa2 = p_sa2.add_run('  ผู้ใช้ในกลุ่มนี้: admin1 (ผู้ดูแลระบบ),  admin3 (กศ.รร.ขส.ขส.ทบ.)')
r_sa2.font.name = FONT_TH
r_sa2.font.size = Pt(13)
r_sa2.font.color.rgb = GRAY_MED
r_sa2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
doc.add_paragraph()

sa_perms = doc.add_table(rows=2, cols=3)
sa_perms.style = 'Table Grid'
sa_cols = [
    ('⚙️  ตั้งค่าระบบ',
     '• ตั้งชื่อโรงเรียน/ชื่อย่อ\n• จัดการรายชื่อครู\n• จัดการรายวิชา\n• ประเภทหลักสูตร/สถานที่'),
    ('👥  จัดการผู้ใช้',
     '• เพิ่ม/ลบ/แก้ไข User\n• กำหนด Role\n• Reset Password\n• ดูประวัติการใช้งาน'),
    ('🌐  อัพเดทเว็บไซท์',
     '• Push Backup ขึ้น GitHub\n• ตั้งค่า GitHub Token\n• Restore ข้อมูลจาก Backup'),
    ('📋  หลักสูตร',
     '• เพิ่ม/ลบ/แก้ไข/คัดลอก\n• เปิด/ปิด Public Mode\n• กำหนดผู้ตรวจถูกต้อง'),
    ('📅  ตารางสอน',
     '• เพิ่ม/แก้ไข/ลบ Entry\n• เพิ่มวันหยุด\n• คัดลอก Entry'),
    ('🖨️  พิมพ์/Export',
     '• พิมพ์ตาราง A4\n• Export Excel\n• บันทึก PDF'),
]
for idx, (title, desc) in enumerate(sa_cols):
    r, c = divmod(idx, 3)
    cell = sa_perms.cell(r, c)
    bg = BLUE_LT if r == 0 else GREEN_LT
    set_cell_bg(cell, bg)
    p_t = cell.paragraphs[0]
    p_t.paragraph_format.space_before = Pt(6)
    p_t.paragraph_format.space_after  = Pt(2)
    r_t = p_t.add_run(title)
    r_t.bold = True
    r_t.font.name = FONT_TH
    r_t.font.size = Pt(12)
    r_t.font.color.rgb = NAVY
    r_t._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    p_d = cell.add_paragraph()
    p_d.paragraph_format.space_before = Pt(0)
    p_d.paragraph_format.space_after  = Pt(6)
    r_d = p_d.add_run(desc)
    r_d.font.name = FONT_TH
    r_d.font.size = Pt(11)
    r_d.font.color.rgb = BLACK
    r_d._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_paragraph()

# --- Admin ---
tbl_adm = doc.add_table(rows=1, cols=1)
tbl_adm.style = 'Table Grid'
cell_adm = tbl_adm.cell(0,0)
set_cell_bg(cell_adm, BLUE_MED)
set_cell_border(cell_adm,
    left={'val':'single','sz':16,'color':'1B6B2E'},
    top={'val':'single','sz':4,'color':'444444'},
    bottom={'val':'single','sz':4,'color':'444444'},
    right={'val':'single','sz':4,'color':'444444'})
p_adm = cell_adm.paragraphs[0]
p_adm.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_adm.paragraph_format.space_before = Pt(8)
p_adm.paragraph_format.space_after  = Pt(4)
r_adm = p_adm.add_run('  🛡️  Admin  —  ผู้ดูแลตารางสอน')
r_adm.bold = True
r_adm.font.name = FONT_TH
r_adm.font.size = Pt(16)
r_adm.font.color.rgb = WHITE
r_adm._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
p_adm2 = cell_adm.add_paragraph()
p_adm2.paragraph_format.space_before = Pt(0)
p_adm2.paragraph_format.space_after  = Pt(8)
r_adm2 = p_adm2.add_run('  ผู้ใช้ในกลุ่มนี้: admin2 (ผตก.รร.ขส.ขส.ทบ.),  num,  fellowship')
r_adm2.font.name = FONT_TH
r_adm2.font.size = Pt(13)
r_adm2.font.color.rgb = GRAY_LT
r_adm2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
doc.add_paragraph()

adm_perms = doc.add_table(rows=1, cols=3)
adm_perms.style = 'Table Grid'
adm_cols = [
    ('📋  หลักสูตร', '✅ เพิ่ม/แก้ไข/ลบ\n✅ คัดลอกหลักสูตร\n❌ ไม่สามารถ Push GitHub\n❌ ไม่สามารถจัดการ User'),
    ('📅  ตารางสอน', '✅ เพิ่ม/แก้ไข/ลบ Entry\n✅ เพิ่มวันหยุด\n✅ บันทึกการแก้ไข (Firebase)\n✅ พิมพ์ตาราง / Export Excel'),
    ('⛔  ข้อจำกัด',  '❌ ไม่สามารถเพิ่ม/ลบ User\n❌ ไม่สามารถ Reset Password\n❌ ไม่เห็นปุ่มอัพเดทเว็บไซท์\n❌ ไม่สามารถตั้งค่าระบบหลัก'),
]
for idx, (title, desc) in enumerate(adm_cols):
    cell = adm_perms.cell(0, idx)
    bgs3 = [GREEN_LT, BLUE_LT, RED_LT]
    set_cell_bg(cell, bgs3[idx])
    p_t = cell.paragraphs[0]
    p_t.paragraph_format.space_before = Pt(6)
    p_t.paragraph_format.space_after  = Pt(2)
    r_t = p_t.add_run(title)
    r_t.bold = True
    r_t.font.name = FONT_TH
    r_t.font.size = Pt(12)
    r_t.font.color.rgb = NAVY
    r_t._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    p_d = cell.add_paragraph()
    p_d.paragraph_format.space_before = Pt(0)
    p_d.paragraph_format.space_after  = Pt(6)
    r_d = p_d.add_run(desc)
    r_d.font.name = FONT_TH
    r_d.font.size = Pt(11)
    r_d.font.color.rgb = BLACK
    r_d._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_paragraph()

# --- Public ---
tbl_pub = doc.add_table(rows=1, cols=1)
tbl_pub.style = 'Table Grid'
cell_pub = tbl_pub.cell(0,0)
set_cell_bg(cell_pub, RGBColor(0x44,0x44,0x44))
set_cell_border(cell_pub,
    left={'val':'single','sz':16,'color':'888888'},
    top={'val':'single','sz':4,'color':'444444'},
    bottom={'val':'single','sz':4,'color':'444444'},
    right={'val':'single','sz':4,'color':'444444'})
p_pub = cell_pub.paragraphs[0]
p_pub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_pub.paragraph_format.space_before = Pt(8)
p_pub.paragraph_format.space_after  = Pt(4)
r_pub = p_pub.add_run('  👁️  Public  —  ผู้เยี่ยมชม (ไม่ต้อง Login)')
r_pub.bold = True
r_pub.font.name = FONT_TH
r_pub.font.size = Pt(16)
r_pub.font.color.rgb = WHITE
r_pub._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
p_pub2 = cell_pub.add_paragraph()
p_pub2.paragraph_format.space_before = Pt(0)
p_pub2.paragraph_format.space_after  = Pt(8)
r_pub2 = p_pub2.add_run('  บุคคลภายนอก นักเรียน ผู้ปกครอง หรือผู้ที่ต้องการดูตารางสอน')
r_pub2.font.name = FONT_TH
r_pub2.font.size = Pt(13)
r_pub2.font.color.rgb = GRAY_MED
r_pub2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
doc.add_paragraph()

pub_perms = doc.add_table(rows=1, cols=2)
pub_perms.style = 'Table Grid'
pub_cols = [
    ('✅  สิ่งที่ทำได้',
     '• ดูตารางสอนหลักสูตรที่เปิด Public\n• พิมพ์ตารางสอน / บันทึก PDF\n• เข้าถึงได้จาก URL โดยตรง\n• ไม่ต้องสมัครสมาชิก'),
    ('❌  สิ่งที่ทำไม่ได้',
     '• ไม่สามารถเพิ่ม/แก้ไข/ลบข้อมูล\n• ไม่เห็นหลักสูตรที่ปิด Public\n• ไม่มีสิทธิ์จัดการ User หรือตั้งค่า\n• ไม่สามารถ Export Excel'),
]
for idx, (title, desc) in enumerate(pub_cols):
    cell = pub_perms.cell(0, idx)
    set_cell_bg(cell, GREEN_LT if idx == 0 else RED_LT)
    p_t = cell.paragraphs[0]
    p_t.paragraph_format.space_before = Pt(6)
    p_t.paragraph_format.space_after  = Pt(2)
    r_t = p_t.add_run(title)
    r_t.bold = True
    r_t.font.name = FONT_TH
    r_t.font.size = Pt(13)
    r_t.font.color.rgb = GREEN if idx == 0 else RED_D
    r_t._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    p_d = cell.add_paragraph()
    p_d.paragraph_format.space_before = Pt(0)
    p_d.paragraph_format.space_after  = Pt(6)
    r_d = p_d.add_run(desc)
    r_d.font.name = FONT_TH
    r_d.font.size = Pt(12)
    r_d.font.color.rgb = BLACK
    r_d._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5: ตาราง User ทั้งหมด
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '05', 'รายชื่อผู้ใช้งานในระบบ', 'System Users')

add_para(doc, 'ผู้ใช้งานที่ลงทะเบียนในระบบ ณ ปัจจุบัน', size=14, space_after=8)

users_data = [
    ('Username',   'ชื่อ-นามสกุล',             'Password',   'Role',        'สิทธิ์'),
    ('admin1',     'ผู้ดูแลระบบ',               'admin1234',  'Super Admin', '👑 สูงสุด'),
    ('admin3',     'กศ.รร.ขส.ขส.ทบ.',           '123456',     'Super Admin', '👑 สูงสุด'),
    ('admin2',     'ผตก.รร.ขส.ขส.ทบ.',          '123456',     'Admin',       '🛡️ ปกติ'),
    ('num',        'num',                        '—',          'Admin',       '🛡️ ปกติ'),
    ('fellowship', 'fellowship',                 '—',          'Admin',       '🛡️ ปกติ'),
]

tbl_users = doc.add_table(rows=len(users_data), cols=5)
tbl_users.style = 'Table Grid'
tbl_users.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row_data in enumerate(users_data):
    cells = tbl_users.row_cells(ri)
    is_hdr = ri == 0
    is_sa  = row_data[3] == 'Super Admin' if not is_hdr else False
    for ci, val in enumerate(row_data):
        cell = cells[ci]
        if is_hdr:
            set_cell_bg(cell, NAVY)
        elif is_sa:
            set_cell_bg(cell, BLUE_LT)
        else:
            set_cell_bg(cell, GRAY_LT if ri % 2 == 1 else WHITE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(val)
        r.bold = is_hdr
        r.font.name = FONT_TH
        r.font.size = Pt(13)
        r.font.color.rgb = WHITE if is_hdr else (NAVY if is_sa else BLACK)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6: การเข้าใช้งาน
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '06', 'วิธีเข้าใช้งาน', 'How to Access')

steps_access = [
    ('🌐', 'เปิด Browser',
     'Chrome, Edge, Firefox, Safari\nรองรับทุกอุปกรณ์ PC / Tablet / Mobile'),
    ('🔗', 'เข้าสู่ URL',
     'https://nattaponatstb.github.io/ats/\nหรือ https://nattaponatstb.github.io/ats/timetable_military.html'),
    ('🔐', 'Login',
     'กรอก Username และ Password\nระบบดึงข้อมูลจาก Firebase อัตโนมัติ'),
    ('📋', 'เลือกหลักสูตร',
     'เลือกหลักสูตรจากแถบด้านซ้าย\nหรือสร้างหลักสูตรใหม่ (เฉพาะ Admin ขึ้นไป)'),
    ('💾', 'บันทึกการแก้ไข',
     'กดปุ่ม "☁️ บันทึกการแก้ไข" เพื่อ Sync\nข้อมูลขึ้น Firebase ให้ผู้อื่นเห็น'),
]

tbl_acc = doc.add_table(rows=len(steps_access), cols=3)
tbl_acc.style = 'Table Grid'
bgs_acc = [NAVY, BLUE_MED, GREEN, ORANGE, RED_D]
for i, (icon, title, desc) in enumerate(steps_access):
    cells = tbl_acc.row_cells(i)
    set_cell_bg(cells[0], bgs_acc[i])
    set_cell_bg(cells[1], BLUE_LT if i % 2 == 0 else WHITE)
    set_cell_bg(cells[2], BLUE_LT if i % 2 == 0 else WHITE)
    p0 = cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    add_run(p0, icon, size=18, color=WHITE)
    p1 = cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(5)
    r1 = p1.add_run(f'ขั้นที่ {i+1}: {title}')
    r1.bold = True
    r1.font.name = FONT_TH
    r1.font.size = Pt(13)
    r1.font.color.rgb = NAVY
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    p2 = cells[2].paragraphs[0]
    p2.paragraph_format.space_before = Pt(5)
    p2.paragraph_format.space_after  = Pt(5)
    r2 = p2.add_run(desc)
    r2.font.name = FONT_TH
    r2.font.size = Pt(12)
    r2.font.color.rgb = BLACK
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)

doc.add_paragraph()

# Note box
tbl_note = doc.add_table(rows=1, cols=1)
tbl_note.style = 'Table Grid'
cell_note = tbl_note.cell(0,0)
set_cell_bg(cell_note, ORANGE_LT)
set_cell_border(cell_note, left={'val':'single','sz':16,'color':'C46A00'})
p_note = cell_note.paragraphs[0]
p_note.paragraph_format.space_before = Pt(8)
p_note.paragraph_format.space_after  = Pt(4)
rn1 = p_note.add_run('  ⚠️  หมายเหตุสำคัญ')
rn1.bold = True
rn1.font.name = FONT_TH
rn1.font.size = Pt(14)
rn1.font.color.rgb = ORANGE
rn1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
for note in [
    'ต้องเชื่อมต่อ Internet เพื่อ Sync ข้อมูลกับ Firebase',
    'ถ้าแก้ไขโดยไม่มี Internet ควรกดปุ่ม "บันทึกการแก้ไข" เมื่อกลับมา Online',
    'Password เก็บในระบบ — ควรเปลี่ยนรหัสผ่าน default ก่อนใช้งานจริง',
]:
    pn = cell_note.add_paragraph()
    pn.paragraph_format.space_before = Pt(1)
    pn.paragraph_format.space_after  = Pt(1)
    rn = pn.add_run(f'      •  {note}')
    rn.font.name = FONT_TH
    rn.font.size = Pt(12)
    rn.font.color.rgb = RGBColor(0x44,0x33,0x00)
    rn._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
p_last = cell_note.add_paragraph()
p_last.paragraph_format.space_after = Pt(8)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), '..', 'docs', 'timetable_system_manual.docx')
# Ensure directory exists
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print('SAVED: ' + out)
