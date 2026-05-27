import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Color Palette Definitions ──
NAVY      = RGBColor(0x1a, 0x2f, 0x4e)
BLUE_MED  = RGBColor(0x2d, 0x5a, 0x8e)
BLUE_LT   = RGBColor(0xd0, 0xe4, 0xf8)
GOLD      = RGBColor(0xc8, 0x9f, 0x2e)
GREEN     = RGBColor(0x1b, 0x6b, 0x2e)
GREEN_LT  = RGBColor(0xd4, 0xed, 0xda)
RED_D     = RGBColor(0x8b, 0x1a, 0x1a)
RED_LT    = RGBColor(0xf8, 0xd7, 0xd7)
GRAY_LT   = RGBColor(0xf4, 0xf6, 0xf9)
GRAY_MED  = RGBColor(0xd0, 0xd8, 0xe4)
WHITE     = RGBColor(0xff, 0xff, 0xff)
BLACK     = RGBColor(0x00, 0x00, 0x00)

FONT_TH = 'TH Sarabun New'

# ── Helper Styling Functions ──
def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(container, text='', size=15, bold=False, italic=False,
             color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    if hasattr(container, 'paragraphs') and hasattr(container, '_tc'):
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p.clear()
    else:
        p = container.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT_TH
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    return p

def add_run(para, text, size=15, bold=False, italic=False, color=BLACK):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_TH
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    return run

def apply_margins(doc):
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 1: โครงการตามแบบฟอร์ม CA002 (CA002_project_proposal_timetable.docx)
# ══════════════════════════════════════════════════════════════════════════════
def build_ca002_proposal(out_path):
    print("Building CA002 project proposal...")
    doc = Document()
    apply_margins(doc)
    
    # Title
    p_title = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p_title, "โครงการพัฒนาระบบจัดการตารางสอนประจำสัปดาห์แบบประสานข้อมูลเรียลไทม์และสำรองข้อมูลผ่านระบบคลาวด์\n", size=18, bold=True)
    add_run(p_title, "กองการศึกษา โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก ประจำปีงบประมาณ 2567", size=15, bold=True)
    
    doc.add_paragraph()
    
    # 1. หลักการและเหตุผล
    add_para(doc, "1. หลักการและเหตุผล", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, 
             "โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก มีภารกิจหลักในการฝึกศึกษาและอบรมทหารกองประจำการ นายสิบนักเรียน และนายทหารนักเรียน "
             "ในหลายหลักสูตรพร้อมกันตลอดปีงบประมาณ การบริหารงานตารางสอนประจำสัปดาห์จึงเป็นฟันเฟืองชิ้นสำคัญในการประสานงานระหว่างกองการศึกษา "
             "แผนกวิชา และผู้สอน ซึ่งในปัจจุบัน ระบบการจัดทำตารางสอนยังเป็นแบบดั้งเดิม (ใช้วิธีพิมพ์กระดาษคำนวณ Excel หรือพิมพ์กระดาษแจกจ่าย) "
             "ทำให้เกิดความล่าช้าในการปรับเปลี่ยนข้อมูลเมื่อมีภารกิจเร่งด่วนของครูผู้สอน หรือเมื่อมีการเปลี่ยนแปลงห้องเรียนกะทันหัน ส่งผลให้เกิดความสับสน "
             "และการสะสมชั่วโมงการสอนของรายวิชาเกิดความคลาดเคลื่อน ส่งผลกระทบโดยตรงต่อคะแนนผลสัมฤทธิ์การประเมินและการประกันคุณภาพการฝึกอบรม",
             size=15)
             
    add_para(doc, 
             "นอกจากนี้ ตามคู่มือการประกันคุณภาพการฝึกอบรมของกองทัพบก พ.ศ. 2566 – 2570 มาตรฐานที่ 1 ตัวชี้วัดที่ 2 ด้านการจัดการเรียนรู้ "
             "และมาตรฐานที่ 2 ตัวชี้วัดที่ 5 ด้านความรู้เทคโนโลยีสารสนเทศของครู/อาจารย์ ระบุถึงความสำคัญของการนำนวัตกรรมและเทคโนโลยีดิจิทัล "
             "มาประยุกต์ใช้เพื่อการบริหารจัดการศึกษา การใช้ตารางสอนรูปแบบเดิมจึงไม่สอดคล้องกับเกณฑ์การประกันคุณภาพฯ และยากต่อการสืบค้นข้อมูลย้อนหลัง "
             "รวมถึงการคำนวณชั่วโมงการสอนเพื่อทำเรื่องเบิกจ่ายค่าสอนตามระเบียบ",
             size=15)

    add_para(doc, 
             "เพื่อการพัฒนาประสิทธิภาพและแก้ปัญหาดังกล่าวอย่างถาวร โครงการนี้จึงเสนอแนวคิดการพัฒนาและใช้นวัตกรรม 'ระบบตารางสอนประจำสัปดาห์แบบเรียลไทม์' "
             "ซึ่งพัฒนาในรูปแบบ Single Page Application (SPA) ทำงานบนเทคโนโลยีคลาวด์และประสานข้อมูลแบบเรียลไทม์ (Real-Time Cloud Database) "
             "ช่วยให้ทุกเครื่องคอมพิวเตอร์และสมาร์ทโฟนของครูสอนและนักเรียนสามารถอัพเดทตารางเรียนที่เปลี่ยนแปลงได้ทันที อีกทั้งยังมีความปลอดภัยข้อมูลสูง "
             "จากการเข้ารหัสผ่านแบบ SHA-256 ช่วยยกระดับการบริหารและอำนวยความสะดวกในการศึกษาของโรงเรียนทหารขนส่งให้ก้าวหน้าอย่างเป็นรูปธรรม",
             size=15)

    # 2. วัตถุประสงค์
    add_para(doc, "2. วัตถุประสงค์", size=16, bold=True, color=NAVY, space_before=12)
    objectives = [
        "เพื่อพัฒนาและติดตั้งระบบการจัดตารางสอนประจำสัปดาห์แบบประสานข้อมูลเรียลไทม์ผ่านคลาวด์ ให้กับโรงเรียนทหารขนส่ง กรมการขนส่งทหารบก",
        "เพื่อเพิ่มความแม่นยำในการคำนวณชั่วโมงเรียนสะสม และตัดปัญหาตารางสอน ครูสอน หรือห้องเรียนซ้ำซ้อนกันในสัปดาห์เรียน",
        "เพื่อยกระดับคะแนนผลการประกันคุณภาพการศึกษาตามมาตรฐานด้านความมั่นคงปลอดภัยข้อมูลและการประยุกต์ใช้เทคโนโลยีดิจิทัล",
    ]
    for idx, obj in enumerate(objectives):
        add_para(doc, f"   2.{idx+1} {obj}", size=15)

    # 3. เป้าหมาย
    add_para(doc, "3. เป้าหมาย", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, "   3.1 เป้าหมายเชิงปริมาณ: กองการศึกษา รร.ขส.ขส.ทบ. สามารถใช้ระบบตารางสอนนี้บริหารจัดการเรียนการสอนครอบคลุมทุกหลักสูตรของโรงเรียนทหารขนส่ง (รองรับการทำตารางสอนพร้อมกันสูงสุด 10 หลักสูตร)", size=15)
    add_para(doc, "   3.2 เป้าหมายเชิงคุณภาพ: ความถูกต้องของข้อมูลชั่วโมงสอนสะสมและตำแหน่งตารางการสอนอยู่ในระดับร้อยละ 100 ปราศจากตารางซ้ำซ้อน และคะแนนความพึงพอใจโดยเฉลี่ยของผู้ใช้งานระบบอยู่ในเกณฑ์ดีมาก (มากกว่าร้อยละ 85)", size=15)

    # 4. วิธีการดำเนินงาน (Gantt Chart Table)
    add_para(doc, "4. วิธีการดำเนินงาน", size=16, bold=True, color=NAVY, space_before=12)
    
    # Gantt Chart Table
    tbl = doc.add_table(rows=8, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    
    headers = ['ลำดับ', 'กิจกรรม/ขั้นตอนการดำเนินงาน', 'ระยะเวลาเริ่มต้น', 'ระยะเวลาสิ้นสุด']
    col_widths = [Cm(1.5), Cm(8.0), Cm(3.2), Cm(3.2)]
    
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=12, bold=True, color=WHITE)
        
    activities = [
        ('1', 'ศึกษารวบรวมข้อมูลปัญหาตารางเรียนและหลักเกณฑ์การประกันคุณภาพฯ', 'มกราคม 2567', 'กุมภาพันธ์ 2567'),
        ('2', 'ออกแบบเค้าโครงระบบ ซอฟต์แวร์ และฐานข้อมูลบน Firebase RTDB', 'มีนาคม 2567', 'มีนาคม 2567'),
        ('3', 'พัฒนาเว็บแอปพลิเคชันเวอร์ชันแรก และระบบความปลอดภัย (SHA-256 Hashing)', 'เมษายน 2567', 'พฤษภาคม 2567'),
        ('4', 'ทดลองใช้งานและตรวจสอบความปลอดภัยการเข้าถึง (Security rules / Secret)', 'มิถุนายน 2567', 'มิถุนายน 2567'),
        ('5', 'ติดตั้งใช้งานจริง ณ กองการศึกษา รร.ขส.ขส.ทบ. และสอนผู้ใช้งานหลัก', 'กรกฎาคม 2567', 'กรกฎาคม 2567'),
        ('6', 'ประเมินผลสัมฤทธิ์และความพึงพอใจการใช้ระบบผ่านแบบสอบถาม', 'สิงหาคม 2567', 'สิงหาคม 2567'),
        ('7', 'สรุปรายงานและเสนอผลโครงการแก่ผู้บังคับบัญชาตามลำดับชั้น', 'กันยายน 2567', 'กันยายน 2567'),
    ]
    
    for row_idx, row_data in enumerate(activities):
        for col_idx, text in enumerate(row_data):
            cell = tbl.cell(row_idx+1, col_idx)
            bg = GRAY_LT if row_idx%2==0 else WHITE
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if col_idx in [0, 2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_run(p, text, size=12, color=BLACK)
            
    doc.add_paragraph()

    # 5. ระยะเวลาการดำเนินโครงการ
    add_para(doc, "5. ระยะเวลาการดำเนินโครงการ", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, "   ตั้งแต่วันที่ 1 มกราคม 2567 ถึงวันที่ 30 กันยายน 2567 (รวมระยะเวลา 9 เดือน)", size=15)

    # 6. งบประมาณ
    add_para(doc, "6. งบประมาณ", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, "   งบประมาณการจัดสร้างและพัฒนาระบบ (โดยจัดสรรจากงบประมาณสนับสนุนการวิจัยศึกษาของหน่วยงาน):", size=15)
    
    tbl_budget = doc.add_table(rows=4, cols=5)
    tbl_budget.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_budget.style = 'Table Grid'
    
    b_headers = ['ลำดับ', 'รายการค่าใช้จ่าย', 'จำนวน : หน่วย', 'ราคา : หน่วย (บาท)', 'รวมทั้งสิ้น (บาท)']
    b_widths = [Cm(1.2), Cm(6.5), Cm(2.5), Cm(2.8), Cm(3.0)]
    
    for i, h in enumerate(b_headers):
        cell = tbl_budget.cell(0, i)
        set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=12, bold=True, color=WHITE)
        
    b_rows = [
        ('1.', 'คู่มือแนะนำระบบ และคู่มือปฏิบัติงานอย่างละเอียด', '5 เล่ม', '200', '1,000'),
        ('2.', 'สื่ออุปกรณ์ทดสอบระบบออฟไลน์ (อุปกรณ์เราเตอร์สำรองและสาย LAN)', '1 ชุด', '2,500', '2,500'),
        ('', 'รวมทั้งสิ้น (สามพันห้าร้อยบาทถ้วน)', '', '', '3,500')
    ]
    
    for row_idx, row_data in enumerate(b_rows):
        cells = tbl_budget.row_cells(row_idx+1)
        is_total = row_idx == 2
        bg = GRAY_LT if row_idx%2==0 else WHITE
        if is_total:
            bg = GOLD
        for col_idx, text in enumerate(row_data):
            cell = cells[col_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if col_idx in [0, 2, 3, 4] and not is_total:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif is_total:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 4 else WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_run(p, text, size=12, bold=is_total, color=WHITE if is_total else BLACK)
            
    doc.add_paragraph()

    # 7. ผู้รับผิดชอบโครงการ
    add_para(doc, "7. ผู้รับผิดชอบโครงการ", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, "   พ.ต.ณัฐพล พัฒนกุล ตำแหน่ง อาจารย์ โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก", size=15)

    # 8. การวัดผลและประเมินผล
    add_para(doc, "8. การวัดผลและประเมินผล", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, "   8.1 ประเมินผลเชิงระบบ: โดยการทดสอบเจาะจงความมั่นคงปลอดภัยฐานข้อมูลและการแฮชรหัสผ่านสำเร็จ 100% รวมถึงระบบออฟไลน์แคชสามารถเก็บตารางได้ถูกต้องครบถ้วน", size=15)
    add_para(doc, "   8.2 ประเมินความพึงพอใจการใช้ระบบ: ทำแบบทดสอบและแบบสอบถามความพอใจจากผู้ใช้ระดับเจ้าหน้าที่แผนกกองการศึกษา และครูอาจารย์ผู้สอน จำนวน 50 นาย โดยเกณฑ์ความพึงพอใจต้องอยู่ในระดับ 'ดีมาก' ขึ้นไป", size=15)

    # 9. สื่อและเอกสารประกอบการประเมินผล
    add_para(doc, "9. สื่อและเอกสารประกอบการประเมินผล", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, "   9.1 แบบทดสอบความปลอดภัยซอฟต์แวร์ (Security Test Report)\n"
                  "   9.2 แบบประเมินผลความพึงพอใจในการใช้ระบบตารางสอนดิจิทัล (แบบมาตราส่วนประมาณค่า 5 ระดับ)\n"
                  "   9.3 รายงานสรุปจำนวนชั่วโมงการสอนจริงรายวิชา", size=15)

    # 10. ผลที่คาดว่าจะได้รับ
    add_para(doc, "10. ผลที่คาดว่าจะได้รับ", size=16, bold=True, color=NAVY, space_before=12)
    add_para(doc, "   10.1 มีระบบซอฟต์แวร์จัดการตารางเรียนเรียลไทม์คลาวด์ใช้งานใน รร.ขส.ขส.ทบ. ทำให้การจัดสรรชั่วโมงครูเป็นระเบียบและไม่มีเหตุการณ์ตารางสอนซับซ้อนชนกัน\n"
                  "   10.2 ลดเวลาและลดการใช้กระดาษในการพิมพ์ตารางสอนอัพเดทใหม่แบบแมนนวลลงได้มากกว่าร้อยละ 90\n"
                  "   10.3 รหัสผ่านผู้ใช้มีความมั่นคงปลอดภัยและไม่ถูกเปิดเผยสู่สาธารณะบนเว็บไซท์โฮสติ้ง ช่วยรักษาความมั่นคงปลอดภัยทางไซเบอร์ของหน่วยงาน", size=15)

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=1, cols=2)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_sig.autofit = False
    
    p_sig1 = tbl_sig.cell(0, 0).paragraphs[0]
    p_sig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_sig1, "ลงชื่อ....................................................ผู้จัดทำโครงการ\n", size=14)
    add_run(p_sig1, "( ณัฐพล พัฒนกุล )\n", size=14, bold=True)
    add_run(p_sig1, "ตำแหน่ง อจ.รร.ขส.ขส.ทบ.", size=13)
    
    p_sig2 = tbl_sig.cell(0, 1).paragraphs[0]
    p_sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_sig2, "ลงชื่อ....................................................ผู้เห็นชอบโครงการ\n", size=14)
    add_run(p_sig2, "( นุรักษ์ ราชรักษ์ )\n", size=14, bold=True)
    add_run(p_sig2, "ตำแหน่ง อจ.หน.แผนกวิชาการขนส่ง รร.ขส.ขส.ทบ.", size=13)

    doc.save(out_path)
    print("SAVED: " + out_path)

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 2: แบบเสนอผลงานนวัตกรรม (01_educational_innovation_proposal_timetable.docx)
# ══════════════════════════════════════════════════════════════════════════════
def build_innovation_proposal(out_path):
    print("Building Educational Innovation Proposal...")
    doc = Document()
    apply_margins(doc)
    
    # ── MEMORANDUM (บันทึกข้อความ) ──
    # Logo (Garuda placeholder symbol via text or just bold centered text)
    p_ตรา = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p_ตรา, "บันทึกข้อความ", size=20, bold=True)
    
    # Memo Details
    tbl_memo = doc.add_table(rows=3, cols=2)
    tbl_memo.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_memo.style = 'Normal Table'
    
    # Row 1
    p_memo00 = tbl_memo.cell(0, 0).paragraphs[0]
    add_run(p_memo00, "ส่วนราชการ  กศ.รร.ขส.ขส.ทบ. (โทร. ๕๒๒๗๒)", size=14, bold=True)
    p_memo01 = tbl_memo.cell(0, 1).paragraphs[0]
    add_run(p_memo01, "โทร. ๕๒๒๗๒", size=14)
    
    # Row 2
    p_memo10 = tbl_memo.cell(1, 0).paragraphs[0]
    add_run(p_memo10, "ที่  กห ๐๔๔๔.๑๓ / ", size=14, bold=True)
    p_memo11 = tbl_memo.cell(1, 1).paragraphs[0]
    add_run(p_memo11, "วันที่   ๒๗   กรกฎาคม   ๒๕๖๗", size=14, bold=True)
    
    # Row 3
    p_memo20 = tbl_memo.cell(2, 0).paragraphs[0]
    add_run(p_memo20, "เรื่อง  ขอส่งผลงานนวัตกรรมด้านการศึกษาเข้าร่วมประกวด ประจำปี ๒๕๖๗", size=14, bold=True)
    
    doc.add_paragraph()
    add_para(doc, "เรียน  ผบ.รร.ขส.ขส.ทบ. (ผ่าน ผตก.รร.ขส.ขส.ทบ.)", size=14, bold=True)
    
    memo_body = (
        "ตามอ้างถึงกองการศึกษา สพ.ศท.ยศ.ทบ. ขอเชิญส่งผลงานวิชาการเข้าร่วมประกวด ประจำปีงบประมาณ ๒๕๖๗ "
        "กองการศึกษา โรงเรียนทหารขนส่ง ขอส่งผลงานนวัตกรรมด้านการศึกษาในหัวข้อ 'ระบบตารางสอนประจำสัปดาห์แบบประสานข้อมูลเรียลไทม์"
        "และสำรองข้อมูลอัตโนมัติ' ซึ่งเป็นสิ่งประดิษฐ์ซอฟต์แวร์ประยุกต์บริหารสถาบันศึกษา พัฒนาขึ้นโดย พ.ต.ณัฐพล พัฒนกุล "
        "โดยได้ดำเนินการลงทะเบียนในระบบเป็นที่เรียบร้อย รายละเอียดตามสิ่งที่ส่งมาด้วยนี้\n\n"
        "จึงเรียนมาเพื่อกรุณาทราบและพิจารณาดำเนินการต่อไป"
    )
    add_para(doc, memo_body, size=14)
    
    p_memo_sig = add_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(p_memo_sig, "พ.อ. ....................................................\n", size=14)
    add_run(p_memo_sig, "( วสันต์ วราสินธุ์ )\n", size=14, bold=True)
    add_run(p_memo_sig, "ผอ.กศ.รร.ขส.ขส.ทบ.", size=13)
    
    doc.add_page_break()
    
    # ── INNOVATION REPORT (แบบเสนอผลงานวิชาการด้านการศึกษา) ──
    p_h2 = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p_h2, "แบบเสนอผลงานวิชาการด้านการศึกษา\n", size=16, bold=True)
    add_run(p_h2, "เข้าร่วมงานประชุมวิชาการและสรุปผลการพัฒนาคุณภาพการศึกษา ประจำปี ๒๕๖๗\n", size=14, bold=True)
    add_run(p_h2, "สำนักการศึกษา กรมยุทธศึกษาทหารบก", size=14, bold=True)
    
    doc.add_paragraph()
    
    # 1. ชื่อผลงาน
    add_para(doc, "1. ชื่อผลงานวิชาการด้านการศึกษา", size=15, bold=True, color=NAVY)
    add_para(doc, "   •  ชื่อภาษาไทย: ระบบจัดการตารางสอนประจำสัปดาห์แบบประสานข้อมูลเรียลไทม์และสำรองข้อมูลอัตโนมัติ\n"
                  "   •  ชื่อภาษาอังกฤษ: Weekly Timetable Management System with Real-Time Cloud Sync & Automated Backup\n"
                  "   •  ประเภทนวัตกรรม: นวัตกรรมสิ่งประดิษฐ์ทางด้านการศึกษา / ระบบบริการจัดการสถานศึกษาดิจิทัล", size=14)
                  
    # 2. รายชื่อคณะผู้จัดทำ
    add_para(doc, "2. รายชื่อคณะผู้จัดทำ", size=15, bold=True, color=NAVY)
    add_para(doc, "   1) พ.ต. ณัฐพล พัฒนกุล (อาจารย์ โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก)\n"
                  "      อีเมล: nattapon.ats01@gmail.com   โทรศัพท์: 090-924-2555", size=14)
                  
    # 3. สถานศึกษา
    add_para(doc, "3. สถานศึกษา/สถาบันการศึกษาที่สังกัด", size=15, bold=True, color=NAVY)
    add_para(doc, "   โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก ค่ายกำแพงเพชรอัครโยธิน", size=14)
    
    # 4. ที่มาของแนวคิดในการพัฒนาผลงาน
    add_para(doc, "4. ที่มาของแนวคิดในการพัฒนาผลงานวิชาการด้านการศึกษา", size=15, bold=True, color=NAVY)
    add_para(doc, 
             "จากการประเมินผลการเรียนการสอนของโรงเรียนทหารขนส่งที่ผ่านมา พบปัญหาหลักในด้านการบริหารจัดการเวลาเรียนและตารางสอน "
             "เนื่องจากหลักสูตรการฝึกศึกษามีหลายรุ่นพร้อมกัน ครูผู้สอนส่วนใหญ่เป็นนายทหารสัญญาบัตรที่มีภารกิจฝึกและปฏิบัติการภายนอกบ่อยครั้ง "
             "ทำให้ตารางเรียนประจำสัปดาห์ต้องปรับปรุงแก้ไขตลอดเวลา เมื่อมีแก้ไขตารางด้วยระบบเดิม (Excel/พิมพ์กระดาษแจก) มักทำให้ข้อมูลตกหล่น "
             "ครูสอนและห้องเรียนเกิดตารางชนกัน และข้อมูลตารางสอนอัพเดทไม่ถึงมือนักเรียนนายสิบในกองร้อยเรียนอย่างทันเวลา "
             "สร้างความเสียหายเชิงการศึกษาและขัดต่อมาตรฐานคุณภาพการฝึกอบรม ยศ.ทบ. พ.ศ. 2566 – 2570 มาตรฐานที่ 1 ตัวชี้วัดที่ 2 "
             "ที่เน้นย้ำเรื่องการนำเทคโนโลยีดิจิทัลมาเพิ่มประสิทธิภาพการเรียนรู้", 
             size=14)
             
    add_para(doc, 
             "ผู้พัฒนาจึงพัฒนา 'ระบบตารางสอนประจำสัปดาห์แบบเรียลไทม์' ขึ้น โดยออกแบบซอฟต์แวร์ให้เก็บข้อมูลบนก้อนคลาวด์ Firebase "
             "และส่งสตรีมข่าวสารแบบเรียลไทม์ข้ามอุปกรณ์ และเพิ่มสิทธิ์รักษาความปลอดภัยรหัสผ่านผู้ดูแลระบบด้วย SHA-256 Hashing "
             "เพื่อป้องกันการสืบข้อมูลความลับส่วนบุคคลตาม พ.ร.บ. คุ้มครองข้อมูลส่วนบุคคล (PDPA) อีกทั้งมีสิทธิ์สำรองข้อมูล data.json "
             "ขึ้นระบบ Git control ของ GitHub Pages อัตโนมัติ ทำให้ระบบตารางสอนมีความเสถียรและทำงานออฟไลน์ได้โดยไม่ต้องเช่าเซิร์ฟเวอร์",
             size=14)

    # 5. วัตถุประสงค์
    add_para(doc, "5. วัตถุประสงค์", size=15, bold=True, color=NAVY)
    add_para(doc, "   1. เพื่อพัฒนาระบบซอฟต์แวร์บริหารตารางสอนเรียลไทม์คลาวด์ ป้องกันตารางชนกันและคำนวณชั่วโมงสะสมวิชาเรียนถูกต้อง 100%\n"
                  "   2. เพื่อยกระดับความมั่นคงปลอดภัยความลับทางทหาร และรักษาความปลอดภัยข้อมูลรหัสผ่านบัญชีแอดมินด้วย SHA-256 Hashing\n"
                  "   3. เพื่อเพิ่มความพึงพอใจการเข้าถึงข้อมูลตารางเรียนครูและนักเรียนผ่านสมาร์ทโฟนได้จากทุกที่ตลอด 24 ชั่วโมง", size=14)

    doc.add_page_break()

    # 6. คุณสมบัติ/คุณลักษณะเฉพาะและขอบเขตการใช้งานของผลงาน
    add_para(doc, "6. คุณสมบัติ/คุณลักษณะเฉพาะและขอบเขตการใช้งานของผลงาน", size=15, bold=True, color=NAVY)
    add_para(doc, "   •  เป็นซอฟต์แวร์ Single Page Application (SPA) โหลดรวดเร็ว ปลอดภัยสูง ทำงานได้บนเว็บเบราว์เซอร์ทุกระบบปฏิบัติการ\n"
                  "   •  การจัดการความปลอดภัย: เข้ารหัสผ่านบัญชีฝั่งเบราว์เซอร์ด้วย SHA-256 และส่ง REST API คู่กับคีย์รักษาความปลอดภัยฐานข้อมูล (?auth=)\n"
                  "   •  ระบบซิงค์เรียลไทม์: เทคโนโลยี Server-Sent Events (SSE) คลาส EventSource ซิงค์ตารางเรียนทันทีเมื่อมีการบันทึกการจัดตารางสอน\n"
                  "   •  ขอบเขตการทำงาน: บริหารชื่อโรงเรียน, ครูอาจารย์, วิชาเรียนตามหลักสูตรสะสม, ห้องเรียน, วันหยุดเรียน และการสั่งพิมพ์ A4 เป็นเลขไทย", size=14)

    # 7. หลักการ วิธีการ และขั้นตอนการทำงานของผลงาน
    add_para(doc, "7. หลักการ วิธีการ และขั้นตอนการทำงานของผลงาน", size=15, bold=True, color=NAVY)
    add_para(doc, "   การพัฒนาดำเนินงานแบ่งเป็นวงจร PDCA (Plan-Do-Check-Action) ตามมาตรฐานพัฒนาคุณภาพการศึกษา ดังนี้:", size=14)
    
    # PDCA Table
    tbl_pdca = doc.add_table(rows=5, cols=2)
    tbl_pdca.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pdca.style = 'Table Grid'
    
    pdca_data = [
        ('ขั้นตอนพัฒนา', 'รายละเอียดกิจกรรม'),
        ('P — วางแผน (Plan)', 'ศึกษาระเบียบชั่วโมงเรียนและคู่มือประกันคุณภาพ ยศ.ทบ. รวบรวมปัญหาตารางซ้อนและครูชนกัน เพื่อเขียนวิเคราะห์โครงสร้างฐานข้อมูล'),
        ('D — ปฏิบัติ (Do)', 'ออกแบบเขียนโปรแกรม HTML/JS/CSS จำลอง localApi และเชื่อมต่อ Firebase Realtime DB, รัน EventSource และพัฒนาความปลอดภัย SHA-256 Hashing'),
        ('C — ตรวจสอบ (Check)', 'จำลองการตั้งค่าระบบ Firebase Rules และคีย์ Secrets เพื่อทดลองแฮกรหัสผ่านแอดมิน และตรวจสอบความมั่นคงปลอดภัยฐานข้อมูล'),
        ('A — ปรับปรุง (Act)', 'นำระบบเข้าทดสอบรันจริง ณ โรงเรียนทหารขนส่ง ในหลักสูตร นนส.ขส. รุ่นที่ 27 และสอนใช้งานพร้อมเขียนบันทึกประวัติข้อเสนอแนะเพื่อพัฒนาต่อยอดแอปพลิเคชันออนไลน์'),
    ]
    
    for row_idx, (p_step, p_desc) in enumerate(pdca_data):
        cells = tbl_pdca.row_cells(row_idx)
        is_hdr = row_idx == 0
        set_cell_bg(cells[0], NAVY if is_hdr else GOLD)
        set_cell_bg(cells[1], GRAY_LT if row_idx%2==1 else WHITE)
        
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        add_run(p0, p_step, size=12, bold=True, color=WHITE)
        
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        add_run(p1, p_desc, size=11, color=BLACK)

    doc.add_paragraph()

    # 8. จุดเด่น หรือกลไกการทำงานที่เป็นจุดเด่นที่แตกต่างจากของผู้อื่น
    add_para(doc, "8. จุดเด่น หรือกลไกการทำงานที่เป็นจุดเด่นที่แตกต่างจากของผู้อื่น", size=15, bold=True, color=NAVY)
    add_para(doc, 
             "   1)  **ไร้ค่าใช้จ่ายเซิร์ฟเวอร์ (Zero-Server-Hosting Cost)**: ระบบออกแบบโดยใช้สถาปัตยกรรมไร้เซิร์ฟเวอร์ พึ่งพา Firebase Free Tier และ GitHub Pages ทำให้หน่วยงานทหารขนส่งไม่มีภาระค่าเช่าเซิร์ฟเวอร์รายปีในการใช้งานแอปพลิเคชัน\n"
             "   2)  **ความปลอดภัยข้อมูลส่วนบุคคลฝั่งไคลเอนต์**: การแฮชรหัสผ่านในเบราว์เซอร์ด้วย SHA-256 มั่นใจได้ว่าไม่มีคีย์หรือรหัสผ่านเป็นตัวอักษรธรรมดารั่วไหลสู่คลาวด์ และปิดกั้นสิทธิ์อ่านเขียน Firebase จากบุคคลภายนอกด้วยระบบ Database Secrets Rules\n"
             "   3)  **ระบบออฟไลน์แคชและสำรองซ้อน**: แม้ไม่มีอินเทอร์เน็ต แอปพลิเคชันยังคงสามารถเปิดตารางเรียนได้และมีข้อมูลผู้ใช้สำรองดึงมาจาก GitHub snapshot data.json ป้องกันความเสียหายกรณีฐานข้อมูลพังได้อย่างมีเสถียรภาพ", size=14)

    # 9. ประโยชน์และคุณค่าของผลงาน
    add_para(doc, "9. ประโยชน์และคุณค่าของผลงาน", size=15, bold=True, color=NAVY)
    add_para(doc, "   •  **ต่อสถาบันการศึกษา**: ได้ซอฟต์แวร์กลางในการจัดเก็บตารางสอน ประหยัดเวลากว่าระบบเดิม 5 เท่า ลดปัญหาครูชนกัน 100% ประหยัดค่ากระดาษและเอกสารนำเสนอ\n"
                  "   •  **ต่อครูผู้สอน**: สามารถเปิดดูตารางเรียนที่อัพเดทได้ทันทีผ่านมือถือ ลดการประสานงานผิดพลาด และตรวจเช็คภาระชั่วโมงสอนได้แม่นยำ\n"
                  "   •  **ต่อนักเรียนทหาร**: ทราบตารางการฝึกศึกษาที่แน่นอน วางแผนเตรียมเครื่องแต่งกายและการฝึกได้ตรงเวลา ช่วยยกระดับประสิทธิภาพการฝึกอบรม", size=14)

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=1, cols=1)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_sig = tbl_sig.cell(0, 0).paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_sig, "ลงชื่อ....................................................ผู้จัดเสนอผลงาน\n", size=14)
    add_run(p_sig, "( ณัฐพล พัฒนกุล )\n", size=14, bold=True)
    add_run(p_sig, "ตำแหน่ง อจ.รร.ขส.ขส.ทบ.\n", size=13)
    add_run(p_sig, "วันที ๒๗ กรกฎาคม พ.ศ. ๒๕๖๗", size=13)

    doc.save(out_path)
    print("SAVED: " + out_path)

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN EXECUTION
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, '..', 'docs')
    os.makedirs(docs_dir, exist_ok=True)
    
    p1 = os.path.join(docs_dir, 'CA002_project_proposal_timetable.docx')
    p2 = os.path.join(docs_dir, '01_educational_innovation_proposal_timetable.docx')
    
    build_ca002_proposal(p1)
    build_innovation_proposal(p2)
    
    print("\nSuccessfully compiled both proposal templates!")
