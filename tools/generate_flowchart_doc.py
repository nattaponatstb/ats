import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Color Palette Definitions ──
NAVY      = RGBColor(0x1a, 0x2f, 0x4e)
BLUE_MED  = RGBColor(0x2d, 0x5a, 0x8e)
BLUE_LT   = RGBColor(0xd0, 0xe4, 0xf8)
GOLD      = RGBColor(0xc8, 0x9f, 0x2e)
GREEN     = RGBColor(0x1b, 0x6b, 0x2e)
GREEN_LT  = RGBColor(0xd4, 0xed, 0xda)
RED_D     = RGBColor(0x8b, 0x1a, 0x1a)
RED_LT    = RGBColor(0xf8, 0xd7, 0xd7)
ORANGE    = RGBColor(0xc4, 0x6a, 0x00)
ORANGE_LT = RGBColor(0xff, 0xf0, 0xcc)
GRAY_LT   = RGBColor(0xf4, 0xf6, 0xf9)
GRAY_MED  = RGBColor(0xd0, 0xd8, 0xe4)
WHITE     = RGBColor(0xff, 0xff, 0xff)
BLACK     = RGBColor(0x00, 0x00, 0x00)

FONT_TH = 'TH Sarabun New'

# ── Helper Styling Functions ──
def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(container, text='', size=15, bold=False, italic=False,
             color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    if hasattr(container, 'paragraphs') and hasattr(container, '_tc'):
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p.clear()
    else:
        p = container.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT_TH
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    return p

def add_run(para, text, size=15, bold=False, italic=False, color=BLACK):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_TH
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    return run

def section_heading(doc, num, title, subtitle=''):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY)
    set_cell_border(cell,
                    top={'val': 'single', 'sz': 6, 'color': 'C89F2E'},
                    bottom={'val': 'single', 'sz': 6, 'color': 'C89F2E'})
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    r1 = p.add_run(f'{num}  {title}')
    r1.bold = True
    r1.font.name = FONT_TH
    r1.font.size = Pt(18)
    r1.font.color.rgb = WHITE
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(subtitle)
        r2.font.name = FONT_TH
        r2.font.size = Pt(12)
        r2.font.color.rgb = GRAY_MED
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    doc.add_paragraph()

def draw_diagram_box(doc, title, lines, color=BLUE_MED, text_color=WHITE):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    set_cell_border(cell,
                    top={'val': 'single', 'sz': 6, 'color': 'C89F2E'},
                    bottom={'val': 'single', 'sz': 6, 'color': 'C89F2E'},
                    left={'val': 'single', 'sz': 6, 'color': 'C89F2E'},
                    right={'val': 'single', 'sz': 6, 'color': 'C89F2E'})
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = FONT_TH
    r.font.size = Pt(14)
    r.font.color.rgb = text_color
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    for l in lines:
        p_l = cell.add_paragraph()
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(4)
        r_l = p_l.add_run(l)
        r_l.font.name = FONT_TH
        r_l.font.size = Pt(11)
        r_l.font.color.rgb = GRAY_LT if text_color == WHITE else BLACK
        r_l._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    doc.add_paragraph()

def add_arrow(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('│\n▼')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = GOLD

def apply_margins(doc):
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN BUILDER
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(base_dir, '..', 'docs', 'programming_layout_diagram.docx')
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    
    doc = Document()
    apply_margins(doc)
    
    # ── Top Gold Bar ──
    tbl_top = doc.add_table(rows=1, cols=1)
    tbl_top.style = 'Table Grid'
    set_cell_bg(tbl_top.cell(0,0), GOLD)
    tbl_top.cell(0,0).paragraphs[0].paragraph_format.space_before = Pt(4)
    tbl_top.cell(0,0).paragraphs[0].paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ── Main Box ──
    tbl_cover = doc.add_table(rows=1, cols=1)
    tbl_cover.style = 'Table Grid'
    cell_cv = tbl_cover.cell(0, 0)
    set_cell_bg(cell_cv, NAVY)
    set_cell_border(cell_cv,
                    top={'val': 'single', 'sz': 12, 'color': 'C89F2E'},
                    bottom={'val': 'single', 'sz': 12, 'color': 'C89F2E'},
                    left={'val': 'single', 'sz': 12, 'color': 'C89F2E'},
                    right={'val': 'single', 'sz': 12, 'color': 'C89F2E'})
    
    p_icon = cell_cv.paragraphs[0]
    p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_icon.paragraph_format.space_before = Pt(25)
    p_icon.paragraph_format.space_after = Pt(10)
    r_icon = p_icon.add_run('⚙️')
    r_icon.font.size = Pt(45)
    
    p_t1 = cell_cv.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_before = Pt(5)
    p_t1.paragraph_format.space_after = Pt(5)
    r_t1 = p_t1.add_run('แผนผังโครงสร้างและการไหลของข้อมูลโปรแกรม')
    r_t1.bold = True
    r_t1.font.name = FONT_TH
    r_t1.font.size = Pt(22)
    r_t1.font.color.rgb = WHITE
    r_t1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    p_t2 = cell_cv.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t2.paragraph_format.space_before = Pt(2)
    p_t2.paragraph_format.space_after = Pt(5)
    r_t2 = p_t2.add_run('Weekly Timetable Data-Flow & Function Maps')
    r_t2.bold = True
    r_t2.font.name = FONT_TH
    r_t2.font.size = Pt(14)
    r_t2.font.color.rgb = GOLD
    r_t2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    p_t3 = cell_cv.add_paragraph()
    p_t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t3.paragraph_format.space_before = Pt(4)
    p_t3.paragraph_format.space_after = Pt(6)
    r_t3 = p_t3.add_run('โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก')
    r_t3.font.name = FONT_TH
    r_t3.font.size = Pt(15)
    r_t3.font.color.rgb = GRAY_MED
    r_t3._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    p_div = cell_cv.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(10)
    r_div = p_div.add_run('─' * 45)
    r_div.font.color.rgb = GOLD
    r_div.font.size = Pt(10)
    
    p_badge = cell_cv.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_before = Pt(2)
    p_badge.paragraph_format.space_after = Pt(25)
    r_b = p_badge.add_run("🗺️ 🛠️ แผนผังระบบโปรแกรม  |  💾 แผนผังโครงสร้างข้อมูล JSON  |  📊 ฟังก์ชันการทำงาน")
    r_b.font.name = FONT_TH
    r_b.font.size = Pt(12)
    r_b.font.color.rgb = GRAY_MED
    r_b._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata info table
    tbl_bot = doc.add_table(rows=1, cols=3)
    tbl_bot.style = 'Table Grid'
    tbl_bot.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ['ประเภทเอกสาร', 'ระดับความสำคัญ', 'ซอฟต์แวร์สนับสนุน']
    values = ['แผนผังซอฟต์แวร์', 'คู่มือนักพัฒนาซอฟต์แวร์', 'Microsoft Word (DOCX)']
    bgs = [BLUE_LT, GREEN_LT, GRAY_LT]
    for idx, (lbl, val, bg) in enumerate(zip(labels, values, bgs)):
        c = tbl_bot.cell(0, idx)
        set_cell_bg(c, bg)
        p_lbl = c.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_before = Pt(4)
        p_lbl.paragraph_format.space_after = Pt(0)
        add_run(p_lbl, lbl, size=11, bold=True, color=NAVY)
        p_val = c.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_val.paragraph_format.space_before = Pt(0)
        p_val.paragraph_format.space_after = Pt(4)
        add_run(p_val, val, size=11, color=BLACK)
        
    doc.add_page_break()
    
    # ── Section 1 ──
    section_heading(doc, '01', 'แผนผังกระบวนการเริ่มต้นระบบ (App Initialization Flow)', 'ลำดับสเต็ปเมื่อผู้ใช้งานเปิดหน้าตารางสอน')
    
    draw_diagram_box(doc, "1. เบราว์เซอร์ดาวน์โหลดหน้าเว็บหลัก (HTML Load)", 
                     ["• รันโครงร่างหน้าต่างตารางสอน", "• เตรียม Event Listeners เช่น ค้นหาคีย์บอร์ด ปุ่มคลิก และ Dialogs"])
    add_arrow(doc)
    
    draw_diagram_box(doc, "2. โหลดข้อมูลแคชจากบราวเซอร์ (Load Local Storage Cache)", 
                     ["• ดึงข้อมูลคีย์เชื่อมต่อเก่า: ttm_firebase_url, ttm_firebase_secret", 
                      "• ดึงข้อมูลแคชตารางสอน: ttm_courses, ttm_entries, ttm_settings"])
    add_arrow(doc)
    
    draw_diagram_box(doc, "3. ตรวจสอบสิทธิ์ผู้ดูแลระบบ (syncUsersBeforeLogin)", 
                     ["• เรียกอ่านรายชื่อแอดมิน: fbGet('/data/users')", 
                      "• หากดึงคลาวด์สำเร็จ: ทำการ Merge ทับ LocalStorage เพื่อรับผู้ใช้ล่าสุด", 
                      "• หากดึงคลาวด์ล้มเหลว (ออฟไลน์): เรียกดึงสำรองข้อมูล GitHub Pages: loadFromGitHub()"])
    add_arrow(doc)
    
    draw_diagram_box(doc, "4. ตรวจสอบสถานะการเชื่อมต่อครั้งแรก (First-Setup check)", 
                     ["• หากมีผู้ใช้ในระบบ: รันหน้าล็อกอิน (loginScreen.classList.remove('hidden'))", 
                      "• หากระบบว่างเปล่า (ไม่มีแอดมินเลย): เข้าสู่โหมดแอดมินคนแรก (showFirstTimeSetupUI()) เพื่อตั้งชื่อ/รหัสผ่านเริ่มต้น"], 
                     color=ORANGE, text_color=WHITE)
                     
    doc.add_page_break()
    
    # ── Section 2 ──
    section_heading(doc, '02', 'แผนผังกระบวนการล็อกอินและความปลอดภัย (Login & Auth Flow)', 'การตรวจสอบรหัสผ่านและการอัพเกรดรหัสผ่าน')
    
    draw_diagram_box(doc, "1. ผู้ใช้ป้อน Username & Password", 
                     ["• ตรวจสอบข้อมูลว่าง หรือความยาวรหัสผ่าน (ต้อง >= 6 หลักสำหรับแอดมินแรก)"], color=GRAY_LT, text_color=BLACK)
    add_arrow(doc)
    
    draw_diagram_box(doc, "2. เปรียบเทียบบัญชีผู้ใช้ (localApi - /auth/login)", 
                     ["• ค้นหาผู้ใช้ที่มี Username ตรงกันในอาเรย์ users ของ LocalStorage", 
                      "• ดึงค่ารหัสผ่านที่ถูกบันทึกมาตรวจสอบ"], color=BLUE_MED, text_color=WHITE)
    add_arrow(doc)
    
    draw_diagram_box(doc, "3. ตรวจสอบประเภทการจัดเก็บรหัสผ่าน (isSha256 Check)", 
                     ["• ตรวจสอบรหัสผ่านที่จัดเก็บในฐานข้อมูลว่ายาว 64 ตัวอักษรฐาน 16 (HEX) หรือไม่", 
                      "• มีผลลัพธ์แยกตามกรณีการเข้ารหัส"], color=GOLD, text_color=WHITE)
    
    # Branching representation inside tables
    tbl_branch = doc.add_table(rows=1, cols=2)
    tbl_branch.style = 'Table Grid'
    tbl_branch.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Left Branch: Hashed
    c_left = tbl_branch.cell(0, 0)
    set_cell_bg(c_left, GREEN_LT)
    set_cell_border(c_left, left={'val':'single','sz':8,'color':'1B6B2E'})
    add_para(c_left, "กรณี 1: รหัสเป็น SHA-256 แล้ว", size=12, bold=True, color=GREEN)
    add_para(c_left, "1) แฮชรหัสผ่านที่ผู้ใช้ป้อนมาด้วย Web Crypto API: hashPassword(body.password)\n"
                     "2) เปรียบเทียบแฮชตรงกัน -> เข้าสู่ระบบสำเร็จ", size=11, color=BLACK)
                     
    # Right Branch: Plaintext
    c_right = tbl_branch.cell(0, 1)
    set_cell_bg(c_right, RED_LT)
    set_cell_border(c_right, left={'val':'single','sz':8,'color':'8B1A1A'})
    add_para(c_right, "กรณี 2: รหัสเป็นตัวอักษรปกติ (Plaintext)", size=12, bold=True, color=RED_D)
    add_para(c_right, "1) เปรียบเทียบรหัสตรงๆ กับที่กรอก\n"
                      "2) หากตรงกัน เข้าสู่ระบบสำเร็จ และแปลงรหัสผ่านเป็นแฮช SHA-256 ทันที\n"
                      "3) บันทึกทับ Local และยิง fbSet ไปอัพเดททับใน Firebase DB คลาวด์กลาง (Auto-Upgrade)", size=11, color=BLACK)
                      
    doc.add_paragraph()
    add_arrow(doc)
    
    draw_diagram_box(doc, "4. เข้าสู่ระบบสำเร็จ (Login Success)", 
                     ["• บันทึก Token: ttm_session เก็บในเบราว์เซอร์", 
                      "• รีเฟรชสิทธิ์สกรีนและ UI: เปิดการทำงานของหน้าแก้ไข หรือเปิดตารางสอนหลักสูตรตามสิทธิ์"], color=NAVY, text_color=WHITE)
                      
    doc.add_page_break()
    
    # ── Section 3 ──
    section_heading(doc, '03', 'แผนผังโครงสร้างข้อมูล (Database JSON Schema Map)', 'โครงสร้างตารางข้อมูล JSON ใน Firebase Realtime Database')
    
    add_para(doc, 'ฐานข้อมูล Firebase จัดเก็บข้อมูลในรูปแบบต้นไม้ JSON โดยจำแนกออกเป็น 4 กิ่งหลัก:', size=14, space_after=10)
    
    # Firebase structure code representation
    tbl_schema = doc.add_table(rows=1, cols=1)
    tbl_schema.style = 'Table Grid'
    tbl_schema.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_sch = tbl_schema.cell(0, 0)
    set_cell_bg(c_sch, GRAY_LT)
    set_cell_border(c_sch, left={'val': 'single', 'sz': 12, 'color': '1B6B2E'})
    p_sch = c_sch.paragraphs[0]
    p_sch.paragraph_format.space_before = Pt(6)
    p_sch.paragraph_format.space_after = Pt(6)
    schema_text = (
        "{\n"
        "  \"data\": {\n"
        "    \"settings\": {\n"
        "      \"schoolName\": \"โรงเรียนทหารขนส่ง\",\n"
        "      \"schoolShort\": \"รร.ขส.ขส.ทบ.\",\n"
        "      \"teachers\": \"[ { \\\"id\\\": \\\"t1\\\", \\\"name\\\": \\\"สมชาย\\\", \\\"rank\\\": \\\"ร.อ.\\\" } ]\",\n"
        "      \"subjects\": \"[ { \\\"id\\\": \\\"s1\\\", \\\"name\\\": \\\"วิชาขับขี่\\\" } ]\"\n"
        "    },\n"
        "    \"courses\": {\n"
        "      \"course_id_123\": {\n"
        "        \"id\": \"course_id_123\",\n"
        "        \"name\": \"หลักสูตรชั้นนายร้อย\",\n"
        "        \"run_number\": \"๕๙\",\n"
        "        \"start_date\": \"2026-05-01\",\n"
        "        \"end_date\": \"2026-08-31\"\n"
        "      }\n"
        "    },\n"
        "    \"entries\": {\n"
        "      \"entry_id_999\": {\n"
        "        \"id\": \"entry_id_999\",\n"
        "        \"course_id\": \"course_id_123\",\n"
        "        \"type\": \"lesson\",\n"
        "        \"date\": \"2026-05-27\",\n"
        "        \"startTime\": \"08:00\",\n"
        "        \"endTime\": \"12:00\",\n"
        "        \"subject\": \"วิชาการขนส่งทางทหาร\",\n"
        "        \"teacherIds\": [\"t1\"]\n"
        "      },\n"
        "      \"entry_id_holiday\": {\n"
        "        \"id\": \"entry_id_holiday\",\n"
        "        \"course_id\": \"course_id_123\",\n"
        "        \"type\": \"holiday\",\n"
        "        \"date\": \"2026-06-03\",\n"
        "        \"label\": \"วันเฉลิมพระชนมพรรษาฯ\"\n"
        "      }\n"
        "    },\n"
        "    \"users\": {\n"
        "      \"admin_id_555\": {\n"
        "        \"id\": \"admin_id_555\",\n"
        "        \"username\": \"admin1\",\n"
        "        \"password\": \"d3c907a... (SHA-256 Hash ความยาว 64 หลัก)\",\n"
        "        \"role\": \"superadmin\",\n"
        "        \"full_name\": \"ผู้ดูแลระบบหลัก\"\n"
        "      }\n"
        "    }\n"
        "  }\n"
        "}"
    )
    r_sch = p_sch.add_run(schema_text)
    r_sch.font.name = 'Consolas'
    r_sch.font.size = Pt(10)
    r_sch.font.color.rgb = BLACK
    
    doc.add_page_break()
    
    # ── Section 4 ──
    section_heading(doc, '04', 'แผนผังการทำงานของระบบเรียลไทม์ (Real-Time SSE Sync Flow)', 'แผนผังการไหลของข้อมูลระหว่างเบราว์เซอร์เครื่องผู้ใช้เมื่อเปลี่ยนตารางสอน')
    
    draw_diagram_box(doc, "เครื่อง Admin แก้ไขข้อมูล (เช่น เปลี่ยนห้องเรียนของวิชา)", 
                     ["1) เรียกฟังก์ชัน api('PUT', '/entries/id', body)", 
                      "2) บันทึกทับ LocalStorage ของ Admin ทันที", 
                      "3) เรียกฟังก์ชัน _fbPushChange() ส่งข้อมูลไปฐานข้อมูล", 
                      "4) ส่งคำขอ PUT ไปยัง HTTPS Firebase: `${FB_DB}/data/entries/id.json?auth=SECRET`"], 
                     color=BLUE_MED, text_color=WHITE)
    
    add_arrow(doc)
    
    draw_diagram_box(doc, "ฐานข้อมูลคลาวด์ Firebase รับการเปลี่ยนแปลง", 
                     ["1) Firebase ตรวจสอบยืนยันตัวตนด้วย Database Secret ใน ?auth=", 
                      "2) เซฟการทับข้อมูล JSON ในกิ่ง /data/entries/id", 
                      "3) ยิงการเปลี่ยนแปลงแบบ Server-Sent Events (SSE) ข่าวสารไปยังขาลูกค้าทั้งหมด"], 
                     color=GOLD, text_color=WHITE)
                     
    add_arrow(doc)
    
    draw_diagram_box(doc, "เครื่องเบราว์เซอร์ของครูและนักเรียนอื่นๆ (Clients)", 
                     ["1) EventSource ได้ยินเหตุการณ์สตรีม 'put' ชนิดอัพเดทเฉพาะจุด", 
                      "2) หากพบว่าเส้นทางคือ /entries/id จะประมวลผลเซฟลง LocalStorage", 
                      "3) เรียกฟังก์ชัน renderScheduleView() ซ้ำแบบ Dynamic", 
                      "4) ตารางตารางเรียนบนจอของนักเรียนเปลี่ยนห้องเรียนใหม่ทันทีเรียลไทม์"], 
                     color=GREEN, text_color=WHITE)
                     
    doc.save(out_path)
    print("SAVED: " + out_path)
