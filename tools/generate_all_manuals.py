import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Color Palette Definitions ──
NAVY      = RGBColor(0x1a, 0x2f, 0x4e)
BLUE_MED  = RGBColor(0x2d, 0x5a, 0x8e)
BLUE_LT   = RGBColor(0xd0, 0xe4, 0xf8)
GOLD      = RGBColor(0xc8, 0x9f, 0x2e)
GREEN     = RGBColor(0x1b, 0x6b, 0x2e)
GREEN_LT  = RGBColor(0xd4, 0xed, 0xda)
RED_D     = RGBColor(0x8b, 0x1a, 0x1a)
RED_LT    = RGBColor(0xf8, 0xd7, 0xd7)
GRAY_LT   = RGBColor(0xf4, 0xf6, 0xf9)
GRAY_MED  = RGBColor(0xd0, 0xd8, 0xe4)
WHITE     = RGBColor(0xff, 0xff, 0xff)
BLACK     = RGBColor(0x00, 0x00, 0x00)

FONT_TH = 'TH Sarabun New'

# ── Helper Styling Functions ──
def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(container, text='', size=15, bold=False, italic=False,
             color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    if hasattr(container, 'paragraphs') and hasattr(container, '_tc'):
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p.clear()
    else:
        p = container.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT_TH
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    return p

def add_run(para, text, size=15, bold=False, italic=False, color=BLACK):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_TH
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    return run

def section_heading(doc, num, title, subtitle=''):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY)
    set_cell_border(cell,
                    top={'val': 'single', 'sz': 6, 'color': 'C89F2E'},
                    bottom={'val': 'single', 'sz': 6, 'color': 'C89F2E'})
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    r1 = p.add_run(f'{num}  {title}')
    r1.bold = True
    r1.font.name = FONT_TH
    r1.font.size = Pt(18)
    r1.font.color.rgb = WHITE
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(subtitle)
        r2.font.name = FONT_TH
        r2.font.size = Pt(12)
        r2.font.color.rgb = GRAY_MED
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    doc.add_paragraph()

def apply_margins(doc):
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

def make_cover_page(doc, title, subtitle, spec_badge):
    apply_margins(doc)
    
    # Top Gold Bar
    tbl_top = doc.add_table(rows=1, cols=1)
    tbl_top.style = 'Table Grid'
    set_cell_bg(tbl_top.cell(0,0), GOLD)
    tbl_top.cell(0,0).paragraphs[0].paragraph_format.space_before = Pt(4)
    tbl_top.cell(0,0).paragraphs[0].paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Main Box
    tbl_cover = doc.add_table(rows=1, cols=1)
    tbl_cover.style = 'Table Grid'
    cell_cv = tbl_cover.cell(0, 0)
    set_cell_bg(cell_cv, NAVY)
    set_cell_border(cell_cv,
                    top={'val': 'single', 'sz': 12, 'color': 'C89F2E'},
                    bottom={'val': 'single', 'sz': 12, 'color': 'C89F2E'},
                    left={'val': 'single', 'sz': 12, 'color': 'C89F2E'},
                    right={'val': 'single', 'sz': 12, 'color': 'C89F2E'})
    
    p_icon = cell_cv.paragraphs[0]
    p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_icon.paragraph_format.space_before = Pt(25)
    p_icon.paragraph_format.space_after = Pt(10)
    r_icon = p_icon.add_run('⚔️')
    r_icon.font.size = Pt(45)
    
    p_t1 = cell_cv.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_before = Pt(5)
    p_t1.paragraph_format.space_after = Pt(5)
    r_t1 = p_t1.add_run(title)
    r_t1.bold = True
    r_t1.font.name = FONT_TH
    r_t1.font.size = Pt(24)
    r_t1.font.color.rgb = WHITE
    r_t1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    p_t2 = cell_cv.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t2.paragraph_format.space_before = Pt(2)
    p_t2.paragraph_format.space_after = Pt(5)
    r_t2 = p_t2.add_run(subtitle)
    r_t2.bold = True
    r_t2.font.name = FONT_TH
    r_t2.font.size = Pt(14)
    r_t2.font.color.rgb = GOLD
    r_t2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    p_t3 = cell_cv.add_paragraph()
    p_t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t3.paragraph_format.space_before = Pt(4)
    p_t3.paragraph_format.space_after = Pt(6)
    r_t3 = p_t3.add_run('โรงเรียนทหารขนส่ง กรมการขนส่งทหารบก')
    r_t3.font.name = FONT_TH
    r_t3.font.size = Pt(15)
    r_t3.font.color.rgb = GRAY_MED
    r_t3._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    p_div = cell_cv.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(10)
    r_div = p_div.add_run('─' * 45)
    r_div.font.color.rgb = GOLD
    r_div.font.size = Pt(10)
    
    p_badge = cell_cv.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_before = Pt(2)
    p_badge.paragraph_format.space_after = Pt(25)
    r_b = p_badge.add_run(spec_badge)
    r_b.font.name = FONT_TH
    r_b.font.size = Pt(13)
    r_b.font.color.rgb = GRAY_MED
    r_b._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TH)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata info table
    tbl_bot = doc.add_table(rows=1, cols=3)
    tbl_bot.style = 'Table Grid'
    tbl_bot.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ['ชั้นความลับ', 'เอกสารอ้างอิง', 'รูปแบบ']
    values = ['ลับภายในหน่วยงาน', 'ระบบความปลอดภัยรุ่นที่ 2', 'Microsoft Word (DOCX)']
    bgs = [BLUE_LT, GREEN_LT, GRAY_LT]
    for idx, (lbl, val, bg) in enumerate(zip(labels, values, bgs)):
        c = tbl_bot.cell(0, idx)
        set_cell_bg(c, bg)
        p_lbl = c.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_before = Pt(4)
        p_lbl.paragraph_format.space_after = Pt(0)
        add_run(p_lbl, lbl, size=11, bold=True, color=NAVY)
        p_val = c.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_val.paragraph_format.space_before = Pt(0)
        p_val.paragraph_format.space_after = Pt(4)
        add_run(p_val, val, size=11, color=BLACK)
        
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 1: ขั้นตอนและกระบวนการเขียนโปรแกรม (detailed_programming_process.docx)
# ══════════════════════════════════════════════════════════════════════════════
def build_programming_process_doc(out_path):
    print("Building detailed_programming_process.docx...")
    doc = Document()
    make_cover_page(doc, 
                    "ขั้นตอนและกระบวนการพัฒนาโปรแกรมโดยละเอียด", 
                    "System Architecture & Detailed Development Steps", 
                    "🌐 Client-Side App  |  🔥 Secure Firebase RTDB  |  📝 Web Crypto API")
    
    # ── Section 1 ──
    section_heading(doc, '01', 'ภาพรวมระบบและสถาปัตยกรรม (System Architecture)', 'สถาปัตยกรรมทางเทคนิคและส่วนประกอบ')
    add_para(doc, 
             'การพัฒนาเว็บแอปพลิเคชันระบบตารางสอนประจำสัปดาห์ (Weekly Timetable System) ได้รับการออกแบบตามแนวทาง '
             'Single Page Application (SPA) ซึ่งระบบการแสดงผล ส่วนควบคุมหน้าจอ และการประมวลผลถูกเขียนไว้ร่วมกันในไฟล์ '
             'timetable_military.html เพียงไฟล์เดียว ทำให้สามารถเปิดรันได้รวดเร็ว ดึงข้อมูลได้ฉับไว และรองรับการทำ Cache ออฟไลน์ผ่าน LocalStorage ในเครื่องผู้ใช้ โดยมีสถาปัตยกรรมแบ่งเป็น 3 ชั้น ดังนี้:',
             size=14, space_after=12)
    
    layers = [
        ('🖥️  1. ส่วนแสดงผลเบราว์เซอร์ (Client Browser Layer)', 
         'ขับเคลื่อนด้วย JavaScript/HTML/CSS ทั้งหมด ทำงานร่วมกันแบบ Virtual Database Controller '
         'โดยมีอ็อบเจ็กต์ state คอยจัดการและเรนเดอร์เนื้อหาผ่านฟังก์ชันหลักเช่น renderScheduleView() และ renderSettingsView() '
         'ร่วมกับการจัดเก็บข้อมูลสำรองไว้ที่ LocalStorage ในรูปของ JSON เพื่อให้ผู้ใช้งานเข้าถึงข้อมูลได้อย่างรวดเร็วแม้การเชื่อมต่อชั่วคราวติดขัด'),
        ('🔥  2. ฐานข้อมูลคลาวด์เรียลไทม์ (Firebase Realtime Database Layer)', 
         'ทำหน้าที่เป็นฐานข้อมูลหลัก (Central Cloud DB) ผ่านบริการ Google Firebase Realtime Database '
         'ซึ่งทำงานเป็น JSON database ไร้เซิร์ฟเวอร์ส่วนกลางของหน่วยงาน ติดต่อผ่าน HTTPS REST API (สำหรับเขียน/อ่านรายจุด) '
         'และการสตรีมเหตุการณ์ SSE (Server-Sent Events) ผ่านคลาส EventSource ใน JavaScript เพื่อให้เกิดการซิงค์ตารางแบบทันทีทันใด (Real-Time Sync)'),
        ('📦  3. ส่วนจัดเก็บโค้ดและสำรองข้อมูลดิบ (GitHub Repository Layer)', 
         'โฮสต์ไฟล์หน้าเว็บผ่าน GitHub Pages เพื่อเปิดบริการลิงก์สาธารณะ นอกจากนี้ สำหรับสิทธิ์ Super Admin '
         'ระบบจะซิงค์ถ่ายโอนข้อมูล snapshot ล่าสุดของตารางสอนในรูปแบบ data.json บันทึกเก็บเป็น Git commit อัตโนมัติ '
         'เพื่อใช้เป็นระบบกู้คืนข้อมูลสำรอง (Backup & Restore Fallback) ในกรณีที่ฐานข้อมูลคลาวด์เกิดข้อผิดพลาด'),
    ]
    
    tbl = doc.add_table(rows=len(layers), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (title, desc) in enumerate(layers):
        cells = tbl.row_cells(idx)
        set_cell_bg(cells[0], NAVY if idx%2==0 else BLUE_MED)
        set_cell_bg(cells[1], GRAY_LT)
        # title
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_before = Pt(5)
        p0.paragraph_format.space_after = Pt(5)
        add_run(p0, title, size=12, bold=True, color=WHITE)
        # desc
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(5)
        p1.paragraph_format.space_after = Pt(5)
        add_run(p1, desc, size=12, color=BLACK)
        
    doc.add_page_break()
    
    # ── Section 2 ──
    section_heading(doc, '02', 'ขั้นตอนและกระบวนการเขียนโปรแกรม (Development Phase)', 'ลำดับขั้นตอนขั้นตอนการสร้างระบบตารางสอน')
    
    phases = [
        ('ขั้นตอนที่ 1', 'ออกแบบโครงสร้างสไตล์ชีตและ UI (CSS Grid & Flexbox Layout)', 
         'สร้างรากฐานสไตล์ชีตโดยไม่พึ่งพา CSS Framework จากภายนอก เพื่อให้โค้ดโหลดเร็วที่สุด จัดวาง Grid เค้าโครงตารางเรียนให้เป็นระเบียบ '
         'รองรับ Media Queries สำหรับการเข้าใช้งานผ่านอุปกรณ์พกพา และสนับสนุน Dark/Light Theme'),
        ('ขั้นตอนที่ 2', 'จำลองฐานข้อมูลเว็บและ API ภายในเบราว์เซอร์ (Virtual localApi & LocalStorage)', 
         'พัฒนาฟังก์ชัน localApi(method, path, body) คอยจำลองตัวจัดการ API และเขียนโครงสร้างจัดเก็บข้อมูลตารางสอน, '
         'รายชื่อครู และวิชาลงสู่ LocalStorage ในบราวเซอร์ เพื่อเอื้อประโยชน์ให้โปรแกรมทำงานออฟไลน์ได้'),
        ('ขั้นตอนที่ 3', 'พัฒนาฟังก์ชันซิงค์เชื่อมต่อฐานข้อมูลคลาวด์ (Firebase Sync & SSE Integration)', 
         'เขียนระบบติดต่อ Firebase RTDB REST API (ฟังก์ชัน fbGet, fbSet, fbDelete) และเปิดโปรโตคอล SSE Stream '
         'ด้วยคำสั่ง EventSource เพื่อคอยฟังกราฟฟิกข้อมูลที่เปลี่ยนแปลงจากเครื่องอื่นๆ และนำมาแปลงเรนเดอร์ทับหน้าจอทันที'),
        ('ขั้นตอนที่ 4', 'สร้างระบบสำรองข้อมูลและเชื่อมต่อ GitHub (GitHub API Integration)', 
         'เขียนฟังก์ชันดึง API ของ GitHub ในกิ่งการจัดส่ง เพื่อคอยเขียนไฟล์ data.json ขึ้นไปเก็บแบบอัตโนมัติ '
         'ทำให้มีระบบบันทึกสำรองและกู้คืน (Backup & Restore) ตลอดเวลาอย่างปลอดภัย'),
        ('ขั้นตอนที่ 5', 'จัดระบบควบคุมสิทธิ์และประวัติการทำงาน (Role-Based Access & Log System)', 
         'แบ่งระดับสิทธิ์ผู้ใช้เป็น Super Admin, Admin, และ Public (ดูได้อย่างเดียว) พร้อมเขียนระบบบันทึก Log '
         'การทำงานเก็บไว้ใน Local และ Firebase เพื่อคอยแทร็กความเคลื่อนไหวการจัดตารางสอน'),
        ('ขั้นตอนที่ 6', 'ยกระดับความปลอดภัยระบบรหัสผ่านและคีย์ (SHA-256 Crypto & Security Token)', 
         'ปรับปรุงระบบโดยเข้ารหัสผ่านฝั่งไคลเอนต์ด้วย SHA-256 ผ่าน Web Crypto API และป้องกันการงัดแงะ Firebase '
         'โดยสนับสนุนช่องใส่ Firebase URL และ Firebase Security Token ในการตั้งค่าเพื่อส่งรหัสผ่าน ?auth= สำหรับ REST API'),
    ]
    
    tbl_ph = doc.add_table(rows=len(phases), cols=3)
    tbl_ph.style = 'Table Grid'
    tbl_ph.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (step_num, step_name, step_desc) in enumerate(phases):
        cells = tbl_ph.row_cells(idx)
        bg_col = GRAY_LT if idx%2==0 else WHITE
        set_cell_bg(cells[0], GOLD)
        set_cell_bg(cells[1], bg_col)
        set_cell_bg(cells[2], bg_col)
        
        # Step Num
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        add_run(p0, step_num, size=12, bold=True, color=WHITE)
        
        # Step Name
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        add_run(p1, step_name, size=12, bold=True, color=NAVY)
        
        # Step Desc
        p2 = cells[2].paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        add_run(p2, step_desc, size=11, color=BLACK)
        
    doc.add_page_break()
    
    # ── Section 3 ──
    section_heading(doc, '03', 'รายละเอียดโค้ดและระบบความปลอดภัย (Security Implementations)', 'ความปลอดภัยระดับรหัสผ่านและช่องทางการเชื่อมโยงข้อมูล')
    
    add_para(doc, '✦  1. การเข้ารหัสด้วยมาตรฐาน SHA-256 Hashing', size=15, bold=True, color=NAVY, space_before=6)
    add_para(doc, 
             'เดิมทีระบบจะจัดเก็บรหัสผ่านบัญชีผู้ใช้เป็นตัวอักษรธรรมดา (Plaintext) ใน Firebase ซึ่งมีความเสี่ยงหากฐานข้อมูลหลุด '
             'เราจึงนำเทคโนโลยี Web Crypto API มาใช้เพื่อแฮชรหัสผ่านแบบฝั่งเบราว์เซอร์ก่อนทำการส่งออกข้อมูลเสมอด้วยฟังก์ชัน SHA-256 '
             'ซึ่งเป็นระบบ Asynchronous คืนค่าเป็นอาเรย์ไบต์และแปลงเป็นสตริง Hex ความยาว 64 ตัวอักษร:',
             size=13, space_after=6)
    
    # Code block inside a table
    tbl_code = doc.add_table(rows=1, cols=1)
    tbl_code.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_code.style = 'Table Grid'
    c_code = tbl_code.cell(0, 0)
    set_cell_bg(c_code, GRAY_LT)
    set_cell_border(c_code, left={'val': 'single', 'sz': 12, 'color': '2D5A8E'})
    p_code = c_code.paragraphs[0]
    p_code.paragraph_format.space_before = Pt(6)
    p_code.paragraph_format.space_after = Pt(6)
    code_text = (
        "async function hashPassword(password) {\n"
        "  if (!password) return '';\n"
        "  const msgUint8 = new TextEncoder().encode(password);\n"
        "  const hashBuffer = await crypto.subtle.digest('SHA-256', msgUint8);\n"
        "  const hashArray = Array.from(new Uint8Array(hashBuffer));\n"
        "  return hashArray.map(b => b.toString(16).padStart(2, '0')).join('');\n"
        "}"
    )
    r_code = p_code.add_run(code_text)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(11)
    r_code.font.color.rgb = BLACK
    
    doc.add_paragraph()
    
    add_para(doc, '✦  2. ระบบปรับปรุงรหัสผ่านรุ่นเก่าอัตโนมัติ (Credentials Auto-Upgrade)', size=15, bold=True, color=NAVY, space_before=6)
    add_para(doc, 
             'เพื่อป้องกันปัญหาผู้ดูแลระบบเดิมล็อคอินเข้าใช้งานไม่ได้หลังจากอัพเกรดโค้ดใหม่ ระบบของเราได้เพิ่มฟังก์ชันอัพเกรดรหัสผ่าน '
             'โดยเมื่อมีการพยายามเข้าสู่ระบบ โปรแกรมจะตรวจความยาวและรูปแบบรหัสผ่านที่จัดเก็บในฐานข้อมูลก่อน:\n'
             '   •  กรณีเป็นรหัสแบบแฮช 64 หลัก (SHA-256) แล้ว: จะนำรหัสผ่านที่ป้อนมาผ่านฟังก์ชันแฮชเพื่อเปรียบเทียบแฮชตรงกัน\n'
             '   •  กรณีเป็นรหัสแบบข้อความปกติ (Plaintext): จะเปรียบเทียบตรงๆ และเมื่อถูกต้อง จะเข้าสู่ระบบพร้อมแฮชรหัสผ่านให้ใหม่ '
             'บันทึกทับ Local และสั่ง fbSet ไปอัพเดททับในฐานข้อมูล Firebase อัตโนมัติทันที',
             size=13, space_after=10)
    
    add_para(doc, '✦  3. การควบคุมสิทธิ์ฐานข้อมูล Firebase ด้วย Token/Secret', size=15, bold=True, color=NAVY, space_before=6)
    add_para(doc, 
             'หากไม่ได้ตั้งค่า Security Rules บน Firebase คนทั่วไปที่เปิดตรวจแกะโค้ดจะสามารถเข้าถึงและลบฐานข้อมูลทั้งหมดได้ '
             'เราจึงนำฟังก์ชันเชื่อมโยง Token/Secret มาใช้ โดยในหน้าตั้งค่า Admin สามารถนำ "Database Secret" ที่คัดลอกมาจาก '
             'Firebase Console มากรอกบันทึกเก็บไว้ใน LocalStorage ในเครื่องตนเอง จากนั้นทุกครั้งที่ระบบเรียกอ่านเขียนฐานข้อมูล '
             'ผ่าน fbGet(), fbSet(), fbDelete() รวมถึง EventSource Stream จะส่งคีย์ ?auth=SECRET แนบไปด้วย ทำให้สามารถล็อก '
             'กฎความปลอดภัย Firebase Realtime Database Rules ให้มีค่าเป็น read/write: auth != null เพื่อความปลอดภัยสูงสุดได้',
             size=13, space_after=10)
    
    doc.save(out_path)
    print("SAVED: " + out_path)

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 2: คู่มือแนะนำโปรแกรม (program_introduction_manual.docx)
# ══════════════════════════════════════════════════════════════════════════════
def build_program_intro_doc(out_path):
    print("Building program_introduction_manual.docx...")
    doc = Document()
    make_cover_page(doc, 
                    "คู่มือแนะนำโปรแกรมระบบตารางสอนประจำสัปดาห์", 
                    "Weekly Timetable Management System Introduction", 
                    "⚔️ โรงเรียนทหารขนส่ง  |  📊 ☁️ ระบบคลาวด์เรียลไทม์  |  🖨️ พิมพ์ตารางสากล")
    
    # ── Section 1 ──
    section_heading(doc, '01', 'หลักการและเหตุผล (Background & Objective)', 'ความเป็นมาและความจำเป็นของโครงการ')
    add_para(doc, 
             'ตารางเรียนและตารางสอนเป็นองค์ประกอบพื้นฐานที่สำคัญยิ่งในการวางแผนจัดการเรียนการสอนของโรงเรียนทหารขนส่ง '
             'กรมการขนส่งทหารบก อย่างไรก็ตาม ในระบบการจัดเก็บและปรับแต่งตารางสอนเดิมมักใช้วิธีบันทึกข้อมูลแบบดั้งเดิมผ่านทาง '
             'กระดาษพิมพ์ หรือไฟล์กระดาษคำนวณ (Excel) ที่ไม่ได้มีส่วนเชื่อมโยงข้อมูลเรียลไทม์ ทำให้ผู้ตรวจวิชา, '
             'ครูอาจารย์ผู้สอน, และนักเรียนทหารแต่ละกองชั้นไม่สามารถอัพเดทตารางเรียนที่เปลี่ยนแปลงกะทันหันร่วมกันได้ทันเวลา '
             'ก่อให้เกิดความคลาดเคลื่อนและส่งผลกระทบโดยตรงต่อประสิทธิผลการฝึกศึกษา',
             size=14, space_after=10)
    
    add_para(doc, 
             'เพื่อขจัดปัญหาดังกล่าว จึงได้ริเริ่มพัฒนา "ระบบจัดการตารางสอนประจำสัปดาห์แบบเรียลไทม์" นี้ขึ้น '
             'เพื่อปรับปรุงขั้นตอนการจัดเตรียมตารางสอน เปลี่ยนมาสู่การจัดเก็บฐานข้อมูลคลาวด์กลาง (Cloud DB) '
             'ที่รองรับการใช้งานพร้อมกันจากอุปกรณ์ทุกประเภท ทำให้ครูวิชารวมถึงผู้บริหารสามารถตรวจสอบชั่วโมงการทำงานสะสม '
             'และอัพเดทตารางเรียนล่าสุดได้อย่างแม่นยำและรวดเร็วสูงสุด',
             size=14, space_after=12)
             
    # ── Section 2 ──
    section_heading(doc, '02', 'คุณลักษณะที่สำคัญของระบบ (Key Capabilities)', 'ระบบทำงานที่ตอบสนองความจำเป็นของสถาบันศึกษา')
    
    features = [
        ('☁️ ฐานข้อมูลคลาวด์เรียลไทม์ (Real-Time Cloud Sync)', 
         'เชื่อมต่อฐานข้อมูล Google Firebase ทำให้ข้อมูลหลักสูตร ตารางเรียน รายชื่อครู และการอัพเดทชั่วโมงสะสมซิงค์ทันทีโดยอัตโนมัติ '
         'ผู้เขียนตารางบน PC อัพเดทข้อมูล นักเรียนเปิดดูทางหน้าจอมือถือจะเห็นตารางเรียนเปลี่ยนใหม่ทันทีโดยไม่ต้องกดรีเฟรชหน้าเว็บ'),
        ('🔒 ระบบความปลอดภัยมาตรฐาน (Security Shield)', 
         'รหัสผ่านผู้ดูแลระบบจัดเก็บอย่างปลอดภัยด้วยการเข้ารหัสผ่าน SHA-256 ฝั่งเบราว์เซอร์ พร้อมระบบตั้งค่าช่องทางการเชื่อมต่อ '
         'Firebase แบบส่วนตัวผ่าน Database Secrets คัดลอกเก็บในเครื่องผู้ดูแลเท่านั้น มั่นใจได้ว่าข้อมูลสำคัญในเซิร์ฟเวอร์จะไม่ถูกดึงหรือลบ'),
        ('🖨️ การสั่งพิมพ์แบบตอบสนอง (Responsive A4 Print)', 
         'ระบบคำนวณความสูงตารางเรียนและแยกหน้ากระดาษ (Page Break) ให้อัตโนมัติในสไตล์ Landscape A4 '
         'โดยมีหัวตารางพร้อมรายชื่อผู้ตรวจและอนุมัติอย่างเป็นทางการในทุกหน้า พร้อมเปลี่ยนรูปแบบเลขสากลเป็นตัวเลขไทยได้อย่างงดงาม'),
        ('📊 สรุปสถิติชั่วโมงการสอน (Excel Report Exports)', 
         'แสดงผลสรุปชั่วโมงที่ใช้ไปเปรียบเทียบกับชั่วโมงรวมตามหลักสูตรจำแนกรายรายวิชาหรือรายผู้สอน '
         'เพื่อใช้ตรวจสอบภาระงานครูและออกรายงานสรุปสำหรับยื่นเสนอผู้บังคับบัญชาผ่านการดาวน์โหลดไฟล์ Excel ได้ทันที'),
    ]
    
    for title, desc in features:
        add_para(doc, f'✦  {title}', size=15, bold=True, color=NAVY, space_before=4, space_after=2)
        add_para(doc, desc, size=13, space_after=8)
        
    doc.add_page_break()
    
    # ── Section 3 ──
    section_heading(doc, '03', 'ประโยชน์ที่คาดว่าจะได้รับ (Expected Outcomes)', 'ความก้าวหน้าและการพัฒนาในระยะยาว')
    
    benefits = [
        'ลดขั้นตอนในการประสานงานตารางสอนระหว่างแผนกวิชากองการศึกษา',
        'ลดความผิดพลาดในการจัดชั่วโมงสอนและห้องเรียนที่อาจซ้ำซ้อนกัน',
        'ตรวจสอบเวลาเรียนสะสมของนักเรียนทหารให้เป็นไปตามเกณฑ์หลักสูตรอย่างแม่นยำ',
        'ยกระดับสู่ระบบบริหารสถานศึกษาระบบดิจิทัล (Digital Education Support)',
        'ลดการสิ้นเปลืองกระดาษพิมพ์สรุปตารางสอนเมื่อมีกรณีสับเปลี่ยนวิชาฉุกเฉิน',
    ]
    
    tbl_b = doc.add_table(rows=len(benefits), cols=1)
    tbl_b.style = 'Table Grid'
    tbl_b.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(benefits):
        cell = tbl_b.cell(idx, 0)
        set_cell_bg(cell, GRAY_LT if idx%2==0 else WHITE)
        set_cell_border(cell, left={'val': 'single', 'sz': 12, 'color': 'C89F2E'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        add_run(p, f'  ✓  {text}', size=13, bold=True, color=BLACK)
        
    doc.save(out_path)
    print("SAVED: " + out_path)

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 3: คู่มือการใช้งานระบบอย่างละเอียด (detailed_user_manual.docx)
# ══════════════════════════════════════════════════════════════════════════════
def build_user_manual_doc(out_path):
    print("Building detailed_user_manual.docx...")
    doc = Document()
    make_cover_page(doc, 
                    "คู่มือการใช้งานระบบตารางสอนอย่างละเอียด", 
                    "System Operation Guide for Admins & Super Admins", 
                    "👑 จัดการข้อมูลพื้นฐาน  |  📅 จัดตารางเรียนและวันหยุด  |  🖨️ ออกรายงานและพิมพ์")
    
    # ── Section 1 ──
    section_heading(doc, '01', 'การตั้งค่าระบบครั้งแรก (First-Time Setup)', 'การเริ่มต้นรันระบบตารางสอนทหารขนส่ง')
    
    add_para(doc, 
             'เมื่อทำการติดตั้งหรือเข้าเชื่อมโยงหน้าเว็บกับฐานข้อมูล Firebase เปล่าเป็นครั้งแรก '
             'หน้าเว็บจะขึ้นหน้าจอเพื่อเปิดสิทธิ์ช่วยเหลือความปลอดภัยพิเศษ:',
             size=14, space_after=8)
             
    steps = [
        'ขั้นตอนการสร้างแอดมินแรก: เข้าสู่หน้าล็อกอิน ระบบจะตรวจไม่พบบัญชีผู้ใช้บนคลาวด์และแสดงข้อความแจ้งเตือนสีแดง "⚠️ ไม่พบผู้ดูแลระบบในคลาวด์"',
        'ให้ทำการพิมพ์ Username และ Password แอดมินใหม่ที่ต้องการในฟิลด์ล็อกอิน (รหัสผ่านต้องยาวอย่างน้อย 6 ตัวอักษร)',
        'กดปุ่ม "+ สร้างบัญชีและเข้าใช้งาน" ระบบจะแฮชรหัสผ่านด้วย SHA-256 และส่งข้อมูลสร้างเป็น Super Admin หลักขึ้นสู่ Firebase และจัดเก็บลงบราวเซอร์ทันที',
        'หลังจากนั้นระบบจะนำเข้าสู่หน้าจอหลักเพื่อให้จัดตารางสอนและบันทึกการตั้งค่าระบบต่อไป',
    ]
    for idx, s in enumerate(steps):
        add_para(doc, f'   {idx+1}. {s}', size=13, space_after=6)
        
    doc.add_paragraph()
    
    # ── Section 2 ──
    section_heading(doc, '02', 'การจัดการข้อมูลพื้นฐาน (Master Data)', 'การตั้งค่าข้อมูลโครงสร้างหลักสูตรและครูวิชา')
    add_para(doc, 
             'การจัดวางตารางสอนต้องการข้อมูลพื้นฐานที่ถูกต้องเพื่อประหยัดเวลาและลดการสะกดคำผิดในรายวัน '
             'ผู้ใช้ระดับ Admin หรือ Super Admin สามารถเข้าไปกรอกบันทึกข้อมูลเหล่านี้ได้ในเมนู "⚙️ ตั้งค่าระบบ":',
             size=14, space_after=10)
    
    masters = [
        ('🏫 ชื่อสถานศึกษา & ชื่อย่อ', 'ระบุชื่อโรงเรียนหลักและชื่อย่อของหน่วยงาน เช่น โรงเรียนทหารขนส่ง กองการศึกษา (ที่จะไปปรากฏในตารางสอนพิมพ์)'),
        ('👤 ครูอาจารย์', 'เพิ่ม/แก้ไขรายชื่อครูสอน พร้อมระบุ ยศ และชื่อ-นามสกุล และเลือกกลุ่มรายวิชาที่ครูสามารถสอนได้'),
        ('📚 รายวิชา', 'เพิ่มรหัสวิชา ชื่อวิชา และระบุจำนวนชั่วโมงรวมของวิชาตามที่กำหนดไว้ในหลักสูตร เพื่อใช้คำนวณชั่วโมงสะสม'),
        ('📍 สถานที่เรียน', 'เพิ่มสถานที่สำหรับใช้เรียน เช่น ห้องเรียน ๑, ลานฝึกขับขี่รถยนต์, กองการศึกษา รร.ขส.ขส.ทบ.'),
        ('📄 หลักฐาน & การแต่งกาย', 'กรอกตัวเลือกข้อมูลหลักฐานเอกสารอ้างอิง และกำหนดเครื่องแบบการแต่งกายในสัปดาห์เรียน'),
    ]
    
    for title, desc in masters:
        add_para(doc, f'   •  {title}', size=14, bold=True, color=NAVY, space_before=4)
        add_para(doc, f'       {desc}', size=13, space_after=6)
        
    doc.add_page_break()
    
    # ── Section 3 ──
    section_heading(doc, '03', 'การจัดตารางสอนและการจัดการวันหยุด (Timetable Operations)', 'การวางแผนรายสัปดาห์ เพิ่ม แก้ไข คัดลอก และวันหยุด')
    add_para(doc, 
             'เมื่อกำหนดข้อมูลพื้นฐานเรียบร้อยแล้ว ให้เลือกหลักสูตรที่แถบเมนูด้านซ้าย และกดปุ่ม "+ เพิ่มรายการ" เพื่อเริ่มเพิ่มวิชาลงตาราง:',
             size=14, space_after=8)
    
    ops = [
        ('1.', 'เพิ่มรายการวิชาสอน (Add Entry)', 
         'ระบุสัปดาห์ที่เรียน (เป็นเลขอารบิก), วันที่เรียน, เลือกเวลาเริ่มต้น-สิ้นสุด, เลือกรายวิชา (ระบบจะดึงครูที่เกี่ยวข้องมาให้เลือก), '
         'เลือกผู้สอนได้สูงสุด 3 คน, กำหนดวิธีสอน, สถานที่, การแต่งกาย, หลักฐาน และหมายเหตุ จากนั้นกดบันทึก'),
        ('2.', 'การคัดลอกรายการสอน (Duplicate Entry)', 
         'หากมีการสอนวิชาเดิมซ้ำๆ หลายๆ วันในสัปดาห์ ให้คลิกปุ่มคัดลอก (📋) ในคอลัมน์ขวาสุดของรายการนั้น ระบบจะสร้างสำเนาข้อมูลเดิม '
         'โดยผู้จัดสามารถแก้ไขแค่วันที่และชั่วโมงให้สอดคล้อง ซึ่งช่วยประหยัดเวลาการคีย์ข้อมูลอย่างมหาศาล'),
        ('3.', 'การจัดการวันหยุดประจำสัปดาห์ (Holiday)', 
         'หากสัปดาห์การเรียนการสอนวันนั้นเป็นวันหยุดราชการหรือหยุดฝึกศึกษา ให้ผู้จัดกดปุ่ม "🚩 วันหยุด" '
         'ระบุวันและพิมพ์ข้อความวันหยุด (เช่น "วันเฉลิมพระชนมพรรษาฯ") ระบบจะทำการปิดและแสดงแถบสีแดงในตารางอย่างสวยงาม '
         'โดยปุ่มแก้ไข ลบ คัดลอก ของวันหยุด จะรวมอยู่ในคอลัมน์ขวาสุดเช่นเดียวกันกับรายการปกติ'),
    ]
    
    for num, name, desc in ops:
        add_para(doc, f'   {num}  {name}', size=14, bold=True, color=NAVY, space_before=4)
        add_para(doc, f'       {desc}', size=13, space_after=8)
        
    doc.add_page_break()
    
    # ── Section 4 ──
    section_heading(doc, '04', 'การสั่งพิมพ์ตารางสอน และดาวน์โหลด Excel', 'วิธีส่งออกเอกสารไปใช้งานภายนอก')
    
    add_para(doc, '✦  1. การพิมพ์ตารางเรียนออกเป็นกระดาษ A4 (Landscape)', size=15, bold=True, color=NAVY, space_before=6)
    add_para(doc, 
             'เพื่อความสะดวกในการส่งเสนอผู้บังคับบัญชาเซ็นอนุมัติ ให้ทำตามขั้นตอนนี้เพื่อจัดรูปแบบงานพิมพ์:\n'
             '   1)  กดเลือกเมนูหลักสูตรที่ต้องการ และเลือกสัปดาห์ที่เรียน\n'
             '   2)  กดปุ่ม "🖨️ พิมพ์ตารางสอนประจำสัปดาห์" ที่มุมขวาบนของหน้าจอ\n'
             '   3)  หน้าต่าง Preview การพิมพ์จะแสดงขึ้น แถบอนุมัติและลายมือชื่อผู้จัดทำจะถูกคำนวณและแสดงท้ายตารางทุกหน้าโดยอัตโนมัติ\n'
             '   4)  ตั้งค่าพิมพ์ของเบราว์เซอร์:\n'
             '        - ปลายทาง (Destination): บันทึกเป็น PDF หรือส่งเครื่องพิมพ์โดยตรง\n'
             '        - แนวการพิมพ์ (Layout): แนวนอน (Landscape)\n'
             '        - ขนาดกระดาษ (Paper Size): A4\n'
             '        - กราฟิกพื้นหลัง (Background Graphics): **ต้องเช็คเครื่องหมายถูก** เพื่อให้สีกระดาษวันและแถบหัวตารางแสดงผล\n'
             '        - ระยะขอบ (Margins): ไม่มีระยะขอบ (None) หรือระยะขอบขั้นต่ำ (Minimum)',
             size=13, space_after=8)
             
    add_para(doc, '✦  2. การออกรายงานเป็นตารางสรุป Excel', size=15, bold=True, color=NAVY, space_before=6)
    add_para(doc, 
             'เพื่ออำนวยความสะดวกในการจัดสรรค่าสอนครูและการตรวจเทียบภาระการสอนกองการศึกษา:\n'
             '   1)  คลิกเข้าเมนู "📊 สรุปชั่วโมงสอน"\n'
             '   2)  ระบบจะแสดงกราฟิกสรุปชั่วโมงรายวิชา ครูสอน และชั่วโมงทั้งหมดจำแนกรายวัน\n'
             '   3)  ผู้ใช้สามารถเลือกดูสรุปตามเงื่อนไข (เช่น ดูเฉพาะวิชาคณิตศาสตร์ หรือครูผู้สอนท่านใดท่านหนึ่ง)\n'
             '   4)  กดปุ่ม "📤 ดาวน์โหลดตาราง Excel" เพื่อบันทึกเป็นไฟล์ .xlsx สำหรับเปิดรันตารางในคอมพิวเตอร์ต่อไป',
             size=13, space_after=10)
             
    doc.save(out_path)
    print("SAVED: " + out_path)

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN EXECUTION
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    # Define outputs relative to tools folder
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, '..', 'docs')
    os.makedirs(docs_dir, exist_ok=True)
    
    p1 = os.path.join(docs_dir, 'detailed_programming_process.docx')
    p2 = os.path.join(docs_dir, 'program_introduction_manual.docx')
    p3 = os.path.join(docs_dir, 'detailed_user_manual.docx')
    
    build_programming_process_doc(p1)
    build_program_intro_doc(p2)
    build_user_manual_doc(p3)
    
    print("\nSuccessfully generated all requested documentation in docs/ directory!")
